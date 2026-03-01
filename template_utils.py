# template_utils.py

import re

import pandas as pd

from pathlib import Path

from docx import Document

from docxtpl import DocxTemplate

import zipfile

import shutil

import os

from concurrent.futures import ThreadPoolExecutor, as_completed

import time

import logging

import traceback

import io

from fuzzywuzzy import process as fuzzy_process



# Importuri opționale

try:

    from odf.opendocument import load

    from odf.text import P

    HAS_ODF = True

except ImportError:

    HAS_ODF = False



try:

    import psutil

    HAS_PSUTIL = True

except ImportError:

    HAS_PSUTIL = False



try:

    from openpyxl import load_workbook

    HAS_OPENPYXL = True

except ImportError:

    HAS_OPENPYXL = False



try:

    import aiofiles

    import asyncio

    HAS_ASYNC = True

except ImportError:

    HAS_ASYNC = False



try:

    from docx2pdf import convert as convert_to_pdf

    HAS_PDF = True

except ImportError:

    HAS_PDF = False



try:

    from docxcompose.composer import Composer

    HAS_MERGE = True

except ImportError:

    HAS_MERGE = False



from audit import audit

from plugin_manager import plugin_manager

from email_utils import send_email_with_attachments



# Logger

logger = logging.getLogger('HRAudit')



# Variabilă globală pentru script-uri (setată din app)

app_scripts = []



# ------------------------------------------------------------

# Extragere placeholder-i

# ------------------------------------------------------------

def extract_placeholders_from_paragraph(paragraph):

    text = paragraph.text

    return set(re.findall(r'{{(.*?)}}', text))



def extract_placeholders_from_table(table):

    placeholders = set()

    for row in table.rows:

        for cell in row.cells:

            for paragraph in cell.paragraphs:

                placeholders.update(extract_placeholders_from_paragraph(paragraph))

            for nested_table in cell.tables:

                placeholders.update(extract_placeholders_from_table(nested_table))

    return placeholders



def extract_placeholders_from_doc(doc_path):

    try:

        doc = Document(doc_path)

        placeholders = set()

        for paragraph in doc.paragraphs:

            placeholders.update(extract_placeholders_from_paragraph(paragraph))

        for table in doc.tables:

            placeholders.update(extract_placeholders_from_table(table))

        for section in doc.sections:

            if section.header:

                for paragraph in section.header.paragraphs:

                    placeholders.update(extract_placeholders_from_paragraph(paragraph))

            if section.footer:

                for paragraph in section.footer.paragraphs:

                    placeholders.update(extract_placeholders_from_paragraph(paragraph))

        return placeholders

    except Exception as e:

        logger.error(f"Eroare la extragerea din doc {doc_path}: {e}")

        return set()



def extract_placeholders_from_odt(odt_path):

    if not HAS_ODF:

        return set()

    try:

        doc = load(odt_path)

        text = ''

        for elem in doc.getElementsByType(P):

            text += str(elem)

        return set(re.findall(r'{{(.*?)}}', text))

    except Exception as e:

        logger.error(f"Eroare ODT {odt_path}: {e}")

        return set()



def extract_placeholders_from_text(text_content):

    return set(re.findall(r'{{(.*?)}}', text_content))



def extract_placeholders_from_html(html_content):

    text = re.sub(r'<[^>]+>', '', html_content)

    return set(re.findall(r'{{(.*?)}}', text))



def extract_placeholders_from_file(file_path):

    ext = Path(file_path).suffix.lower()

    if ext == '.docx':

        return extract_placeholders_from_doc(file_path)

    elif ext == '.odt' and HAS_ODF:

        return extract_placeholders_from_odt(file_path)

    elif ext in ['.txt', '.csv']:

        with open(file_path, 'r', encoding='utf-8') as f:

            return extract_placeholders_from_text(f.read())

    elif ext in ['.html', '.htm']:

        with open(file_path, 'r', encoding='utf-8') as f:

            return extract_placeholders_from_html(f.read())

    else:

        logger.warning(f"Format neacceptat pentru extragere: {ext}")

        return set()



def extract_all_placeholders_from_files(file_paths, log_queue=None, progress_callback=None):

    all_placeholders = set()

    placeholder_map = {}

    if not file_paths:

        if log_queue:

            log_queue.put("Nu s-au selectat fișiere.")

        return [], {}

    for i, file_path in enumerate(file_paths):

        msg = f"Procesez fișierul {i+1}/{len(file_paths)}: {Path(file_path).name}"

        if log_queue:

            log_queue.put(msg)

        logger.info(msg)

        if progress_callback:

            progress_callback(i+1, len(file_paths))

        try:

            placeholders = extract_placeholders_from_file(file_path)

            all_placeholders.update(placeholders)

            for ph in placeholders:

                if ph not in placeholder_map:

                    placeholder_map[ph] = []

                placeholder_map[ph].append(file_path)

        except Exception as e:

            logger.error(f"Eroare la procesarea fișierului {file_path}: {e}")

            if log_queue:

                log_queue.put(f"Eroare la procesarea {file_path}: {e}")

    return sorted(all_placeholders), placeholder_map



async def extract_placeholders_from_file_async(file_path):

    ext = Path(file_path).suffix.lower()

    if ext == '.docx':

        loop = asyncio.get_event_loop()

        return await loop.run_in_executor(None, extract_placeholders_from_doc, file_path)

    elif ext == '.odt' and HAS_ODF:

        loop = asyncio.get_event_loop()

        return await loop.run_in_executor(None, extract_placeholders_from_odt, file_path)

    elif ext in ['.txt', '.html', '.htm']:

        if HAS_ASYNC:

            async with aiofiles.open(file_path, 'r', encoding='utf-8') as f:

                content = await f.read()

        else:

            with open(file_path, 'r', encoding='utf-8') as f:

                content = f.read()

        if ext == '.txt':

            return extract_placeholders_from_text(content)

        else:

            return extract_placeholders_from_html(content)

    else:

        return set()



async def extract_all_placeholders_async(file_paths, progress_callback=None):

    tasks = []

    for i, fp in enumerate(file_paths):

        tasks.append(extract_placeholders_from_file_async(fp))

    results = await asyncio.gather(*tasks)

    all_placeholders = set()

    placeholder_map = {}

    for i, placeholders in enumerate(results):

        for ph in placeholders:

            all_placeholders.add(ph)

            if ph not in placeholder_map:

                placeholder_map[ph] = []

            placeholder_map[ph].append(file_paths[i])

        if progress_callback:

            progress_callback(i+1, len(file_paths))

    return sorted(all_placeholders), placeholder_map



def generate_excel_template(placeholders, output_excel, log_queue=None):

    try:

        df = pd.DataFrame(columns=['ID'] + placeholders)

        df.to_excel(output_excel, index=False)

        if log_queue:

            log_queue.put(f"Fișier Excel generat: {output_excel}")

            log_queue.put("Completați acest fișier cu date (fiecare rând = un set de valori).")

        logger.info(f"Excel generat cu {len(placeholders)} coloane: {output_excel}")

        return placeholders

    except Exception as e:

        logger.error(f"Eroare la generarea Excel: {e}")

        if log_queue:

            log_queue.put(f"Eroare la generarea Excel: {e}")

        return []



# ------------------------------------------------------------

# Randare documente

# ------------------------------------------------------------

def render_document_from_template(template_path, context, output_path):

    ext = Path(template_path).suffix.lower()

    

    # Detecție simplă a tuturor placeholder-urilor din fisier inainte de a randare

    all_placeholders_in_file = extract_placeholders_from_file(template_path)

    found_placeholders = [p for p in all_placeholders_in_file if p in context and str(context[p]).strip() != '']

    missing_placeholders = [p for p in all_placeholders_in_file if p not in context or str(context[p]).strip() == '']



    if ext == '.docx':

        doc = DocxTemplate(template_path)

        doc.render(context)

        doc.save(output_path)

    elif ext == '.odt' and HAS_ODF:

        doc = load(template_path)

        for elem in doc.getElementsByType(P):

            if elem.firstChild and elem.firstChild.nodeType == elem.TEXT_NODE:

                text = elem.firstChild.data

                for key, val in context.items():

                    text = text.replace('{{' + key + '}}', str(val))

                elem.firstChild.data = text

        doc.save(output_path)

    elif ext in ['.txt', '.html', '.htm']:

        with open(template_path, 'r', encoding='utf-8') as f:

            content = f.read()

        for key, val in context.items():

            content = content.replace('{{' + key + '}}', str(val))

        with open(output_path, 'w', encoding='utf-8') as f:

            f.write(content)

    else:

        raise ValueError(f"Format neacceptat: {ext}")

        

    return found_placeholders, missing_placeholders



def iter_excel_chunks(file_path, sheet_name=0, chunksize=1000):

    """Read Excel in chunks using pandas (reliable for all xlsx variants including shared strings)."""

    print(f"DEBUG EXCEL: reading {file_path} sheet={sheet_name} chunksize={chunksize}", flush=True)

    try:

        # Use pandas to read the full sheet then yield in chunks

        # This avoids the openpyxl read_only=True mode bug with shared strings

        df_full = pd.read_excel(file_path, sheet_name=sheet_name, engine='openpyxl')

        print(f"DEBUG EXCEL: pandas read {len(df_full)} rows, columns={list(df_full.columns)}", flush=True)

        if len(df_full) == 0:

            return

        for start in range(0, len(df_full), chunksize):

            chunk = df_full.iloc[start:start + chunksize].copy()

            if HAS_PSUTIL:

                mem = psutil.virtual_memory()

                if mem.available < 500 * 1024 * 1024:

                    logger.warning("Memorie redusă în timpul procesării chunk-urilor Excel")

            print(f"DEBUG EXCEL: yielding chunk rows {start}-{start+len(chunk)}", flush=True)

            yield chunk

    except Exception as e:

        logger.error(f"Eroare la citirea Excel în chunks: {e}")

        print(f"DEBUG EXCEL ERROR: {e}", flush=True)

        raise



def clean_dataframe(df):

    """Curăță datele: elimină spații, title case pentru nume, etc."""

    df = df.copy()

    # Safely strip whitespace from object columns (may contain mixed types)

    for col in df.select_dtypes(include=['object']).columns:

        try:

            df[col] = df[col].astype(str).str.strip().replace('nan', '')

        except Exception:

            pass  # skip non-string-convertible columns

    # Title case pentru coloane care conțin 'nume' sau 'prenume'

    for col in df.columns:

        if 'nume' in col.lower() or 'prenume' in col.lower() or 'name' in col.lower():

            try:

                df[col] = df[col].astype(str).str.title().replace('Nan', '')

            except Exception:

                pass

    return df





def fuzzy_match_columns(placeholders, df_columns, threshold=80):

    """

    Sugerează corespondențe între placeholder-ele din șablon și coloanele Excel.

    Returnează un dicționar: {placeholder: cea_mai_bună_coloană}

    """

    matches = {}

    for p in placeholders:

        # Încearcă potrivire exactă (case-insensitive)

        exact = next((c for c in df_columns if c.lower() == p.lower()), None)

        if exact:

            matches[p] = exact

            continue

        

        # Încearcă potrivire fuzzy

        best_match, score = fuzzy_process.extractOne(p, df_columns)

        if score >= threshold:

            matches[p] = best_match

    return matches



def render_documents(template_files, data_file, output_folder, folder_column, sheets=None, log_queue=None, resume_from=0,

                     progress_callback=None, stop_event=None, filename_pattern="{ID}.docx", parallel=False, chunksize=0,

                     email_config=None, send_mode='individual', email_column=None, email_subject_pattern=None,

                     pdf_gen=False, zip_gen=False, merge_gen=False, audio_alert=False, subfolder_col="(Niciunul)", clean_data=False,

                     zip_per_row=False):

    # Convertim output_folder la Path

    if output_folder is None:

        output_folder = Path("output")

    else:

        output_folder = Path(output_folder)

    output_folder.mkdir(exist_ok=True, parents=True)



    # Colectăm toate placeholder-ele din toate șabloanele pentru fuzzy mapping

    all_placeholders = set()

    for tf in template_files:

        all_placeholders.update(extract_placeholders_from_file(tf))

    

    fuzzy_map = {} # va fi populat după încărcarea Excel-ului



    # Încărcăm foile

    try:

        xl = pd.ExcelFile(data_file)

        all_sheets = xl.sheet_names

        if sheets is None or sheets == [None]:

            sheets = all_sheets

    except Exception as e:

        error_msg = f"Eroare la citirea fișierului Excel: {str(e)}"

        logger.error(error_msg)

        if log_queue:

            log_queue.put(error_msg)

        return



    total_processed = 0

    all_generated_files = []  # pentru ZIP și Merge



    for sheet in sheets:

        if stop_event and stop_event.is_set():

            break

        # Determinare index/nume foaie pentru openpyxl

        if isinstance(sheet, int):

            sheet_idx = sheet  # 0-based index

            try:

                xl_temp = pd.ExcelFile(data_file)

                sheet_name_str = xl_temp.sheet_names[sheet_idx] if sheet_idx < len(xl_temp.sheet_names) else xl_temp.sheet_names[0]

            except:

                sheet_name_str = 0

        else:

            sheet_name_str = sheet

        if log_queue:

            log_queue.put(f"Procesez foaia: {sheet_name_str}")

        sheet_folder = output_folder

        sheet_folder.mkdir(exist_ok=True)



        try:

            if chunksize > 0:

                try:

                    wb_count = load_workbook(data_file, read_only=True)

                    ws_count = wb_count[sheet_name_str]

                    total_rows = ws_count.max_row - 1

                    wb_count.close()

                except:

                    total_rows = 0

                chunk_iter = iter_excel_chunks(data_file, sheet_name=sheet_name_str, chunksize=chunksize)

            else:

                df = pd.read_excel(data_file, sheet_name=sheet if isinstance(sheet, int) else sheet_name_str)

                total_rows = len(df)

                chunk_iter = [df]

        except Exception as e:

            error_msg = f"Eroare la citirea foii {sheet}: {str(e)}"

            logger.error(error_msg)

            if log_queue:

                log_queue.put(error_msg)

            continue



        # Determinăm datele pentru "first_chunk" fără a consuma iteratorul

        if chunksize == 0:

            if clean_data:

                df = clean_dataframe(df)

                chunk_iter = [df]

            first_chunk = df

        else:

            first_chunk = pd.read_excel(data_file, sheet_name=sheet if isinstance(sheet, int) else sheet_name_str, nrows=1)

            chunk_iter = iter_excel_chunks(data_file, sheet_name=sheet_name_str, chunksize=chunksize)

            if clean_data:

                chunk_iter = (clean_dataframe(chunk) for chunk in chunk_iter)



        use_column = folder_column and folder_column in first_chunk.columns

        if not use_column and folder_column:

            if log_queue:

                log_queue.put(f"Coloana '{folder_column}' nu există. Se va folosi 'ID' pentru numele folderelor.")

            logger.warning(f"Coloana '{folder_column}' nu există în Excel.")



        # Calculăm fuzzy mapping o singură dată per generare

        if log_queue:

            log_queue.put("Se calculează maparea inteligentă a coloanelor...")

        fuzzy_map = fuzzy_match_columns(all_placeholders, first_chunk.columns.tolist())

        if log_queue and fuzzy_map:

            log_queue.put(f"✅ Mapare inteligentă finalizată ({len(fuzzy_map)} corespondențe găsite).")



        checkpoint_file = sheet_folder / "checkpoint.txt"

        start_idx = resume_from if sheet == sheets[0] else 0

        if resume_from == 0 and checkpoint_file.exists():

            try:

                with open(checkpoint_file, 'r') as f:

                    saved = int(f.read().strip())

                    if saved < total_rows:

                        start_idx = saved

                        if log_queue:

                            log_queue.put(f"Se reia de la rândul {start_idx+1} în foaia {sheet} (checkpoint găsit).")

            except:

                pass



        global app_scripts

        processed = 0

        total_to_process = total_rows - start_idx



        def process_row(row_data, row_idx):

            start_time = time.time()

            if 'ID' not in row_data:

                row_data['ID'] = row_idx + 1

            row_id = row_data['ID']

            if use_column:

                val = row_data[folder_column]

                if pd.isna(val) or str(val).lower() == 'nan':

                    folder_name = f"rand_{row_id}"

                else:

                    folder_name = str(val).strip()

                    folder_name = re.sub(r'[<>:"/\\|?*\n\r\t]', '_', folder_name)

                    if not folder_name:

                        folder_name = f"rand_{row_id}"

            else:

                folder_name = f"rand_{row_id}"



            context = {k: ('' if pd.isna(v) else v) for k, v in row_data.items()}

            

            # Aplicăm fuzzy mapping pentru câmpurile care lipsesc

            for ph, col in fuzzy_map.items():

                if ph not in context or str(context[ph]).strip() == '':

                    context[ph] = context.get(col, '')



            for script in app_scripts:

                try:

                    context = script(context)

                except Exception as e:

                    logger.error(f"Eroare în script personalizat: {e}")



            # Dacă avem și subfolder_col (vechi), îl putem combina, 

            # dar prioritatea este folder_name (din folder_column)

            final_folder = sheet_folder / folder_name

            if subfolder_col and subfolder_col != "(Niciunul)" and subfolder_col in context:

                sub_val = context[subfolder_col]

                if not pd.isna(sub_val) and str(sub_val).lower() != 'nan':

                    sub = str(sub_val).strip()

                    sub = re.sub(r'[<>:"/\\|?*\n\r\t]', '_', sub)

                    if sub:

                        final_folder = final_folder / sub

            

            row_folder = final_folder

            row_folder.mkdir(exist_ok=True, parents=True)



            # Colectăm căile fișierelor generate pentru acest rând (pentru email, zip, merge)

            generated_files = []



            for template_path in template_files:

                try:

                    template_stem = Path(template_path).stem

                    context['template_name'] = template_stem



                    # Generăm numele fișierului

                    fname = filename_pattern

                    for key, val in context.items():

                        fname = fname.replace('{' + key + '}', str(val))

                    # Curățăm caracterele invalide și whitespace distructiv

                    fname = re.sub(r'[<>:"/\\|?*\n\r\t]', '_', fname)

                    # Limităm lungimea numelui pentru a evita erori de sistem (Windows MAX_PATH)

                    if len(fname) > 150:

                        fname = fname[:150]

                    

                    ext = Path(template_path).suffix

                    if not fname.endswith(ext):

                        fname += ext

                    output_path = row_folder / fname

                    found_ph, missing_ph = render_document_from_template(template_path, context, output_path)

                    generated_files.append(output_path)

                    all_generated_files.append(output_path)



                    # Pregătim date pentru audit, incluzând ce s-a înlocuit și ce lipsește

                    audit_entry = {

                        'row_id': row_id,

                        'file': template_stem,

                        'output': fname,

                        'status': 'success',

                        'found_placeholders': found_ph,

                        'missing_placeholders': missing_ph

                    }

                    if 'Nume' in context:

                        audit_entry['nume'] = context['Nume']

                    audit_entry['batch_render'] = True

                    audit.log(action='render', details=audit_entry)

                    if log_queue:

                        log_queue.put(('SUCCESS', row_id))



                    # Export plugin

                    for fmt in plugin_manager.get_export_formats():

                        out_export = row_folder / f"{Path(output_path).stem}.{fmt}"

                        plugin_manager.export_document(str(output_path), fmt, str(out_export))



                    # PDF generation

                    if pdf_gen and HAS_PDF:

                        try:

                            pdf_path = output_path.with_suffix('.pdf')

                            convert_to_pdf(str(output_path), str(pdf_path))

                            generated_files.append(pdf_path)

                            all_generated_files.append(pdf_path)

                            if log_queue:

                                log_queue.put(f"   PDF generat: {pdf_path.name}")

                        except Exception as e:

                            if log_queue:

                                log_queue.put(f"   ⚠Eroare PDF: {e}")



                    # Email individual

                    if send_mode != 'none' and email_config and email_config.get('enabled') and email_column and email_column in context:

                        pass  # se va face după bucla template



                except Exception as e:

                    error_msg = f"Eroare la generarea documentului {Path(template_path).name} pentru rândul {row_id} în foaia {sheet}: {str(e)}"

                    logger.error(f"{error_msg}\n{traceback.format_exc()}")

                    if log_queue:

                        log_queue.put(error_msg)

                    

                    audit_entry = {

                        'row_id': row_id,

                        'file': Path(template_path).stem,

                        'status': 'error',

                        'error': str(e),

                        'found_placeholders': [],

                        'missing_placeholders': []

                    }

                    audit_entry['batch_render'] = True

                    audit.log(action='render_error', details=audit_entry)

                    if log_queue:

                        log_queue.put(('ERROR', row_id))



            # Trimitere email individual sau row_zip dacă este activat

            if send_mode in ['individual', 'row_zip'] and email_config and email_config.get('enabled') and email_column and email_column in context:

                try:

                    recipient_email = context[email_column]

                    if not recipient_email or '@' not in str(recipient_email):

                        if log_queue:

                            log_queue.put(f"⚠Adresă email invalidă pentru rândul {row_id}: '{recipient_email}' - se omite trimiterea.")

                    else:

                        subject = email_config['subject']

                        body = email_config['body']

                        for key, val in context.items():

                            subject = subject.replace('{' + key + '}', str(val))

                            body = body.replace('{' + key + '}', str(val))

                        if email_subject_pattern:

                            subject = email_subject_pattern

                            for key, val in context.items():

                                subject = subject.replace('{' + key + '}', str(val))



                        email_cfg = {

                            'smtp_server': email_config['smtp_server'],

                            'smtp_port': email_config['smtp_port'],

                            'username': email_config['username'],

                            'password': email_config['password'],

                            'from': email_config['from'],

                            'to': recipient_email,

                            'subject': subject,

                            'body': body,

                            'no_auth': email_config.get('no_auth', False)

                        }

                        

                        files_to_send = generated_files

                        if send_mode == 'row_zip' and generated_files:

                            # Creăm o arhivă zip specifică rândului

                            row_zip_path = row_folder / f"Documente_Rand_{row_id}.zip"

                            with zipfile.ZipFile(row_zip_path, 'w', zipfile.ZIP_DEFLATED) as rzip:

                                for gf in generated_files:

                                    if os.path.exists(gf):

                                        rzip.write(gf, Path(gf).name)

                            files_to_send = [row_zip_path]

                            if log_queue:

                                log_queue.put(f"   📦 Arhivă parțială creată pentru rândul {row_id}: {row_zip_path.name}")

                        

                        if log_queue:

                            log_queue.put(f"📧 Trimite email către {recipient_email} cu subiectul: {subject}")

                        send_email_with_attachments(email_cfg, files_to_send, log_queue, context)

                        time.sleep(1)

                except Exception as e:

                    error_msg = f"Eroare la trimiterea emailului pentru rândul {row_id}: {e}"

                    logger.error(error_msg)

                    if log_queue:

                        log_queue.put(f"❌ {error_msg}")



            # Arhivă ZIP per rând dacă este activată

            if zip_per_row and generated_files:

                try:

                    row_zip_name = f"Arhiva_{folder_name}.zip"

                    # Punem ZIP-ul în folderul părinte al row_folder sau chiar în el

                    # User-ul a zis: "arhiva zip cu toate documentele dintr-un folder, arhiva ar trebui sa fie in dosar pentru documentele generate/cate o arhiva pentru fiecare folder"

                    # Deci în interiorul row_folder? Sau lângă?

                    # "cate o arhiva pentru fiecare folder" - probabil în interiorul lui

                    row_zip_path = row_folder / row_zip_name

                    with zipfile.ZipFile(row_zip_path, 'w', zipfile.ZIP_DEFLATED) as rzip:

                        for gf in generated_files:

                            if os.path.exists(gf):

                                rzip.write(gf, Path(gf).name)

                    if log_queue:

                        log_queue.put(f"   📦 Arhivă ZIP creată pentru {folder_name}: {row_zip_name}")

                except Exception as e:

                    if log_queue:

                        log_queue.put(f"   ⚠Eroare creare arhivă rând: {e}")



            elapsed = time.time() - start_time

            if log_queue:

                log_queue.put(f"Profilare: rând {row_id} procesat în {elapsed:.2f} secunde")

            return row_idx



        # Procesare paralelă sau secvențială

        if parallel:

            with ThreadPoolExecutor(max_workers=4) as executor:

                futures = []

                current_global_idx = start_idx

                for chunk in chunk_iter:

                    for local_idx, row in chunk.iterrows():

                        if stop_event and stop_event.is_set():

                            break

                        future = executor.submit(process_row, row.to_dict(), current_global_idx)

                        futures.append(future)

                        current_global_idx += 1

                    if stop_event and stop_event.is_set():

                        break

                for i, future in enumerate(as_completed(futures)):

                    if stop_event and stop_event.is_set():

                        executor.shutdown(wait=False)

                        break

                    processed += 1

                    if progress_callback:

                        progress_callback(total_processed + processed, total_to_process + (total_processed if total_processed>0 else 0))

                    if processed % 10 == 0:

                        with open(checkpoint_file, 'w') as f:

                            f.write(str(start_idx + processed))

        else:

            for chunk in chunk_iter:

                print(f"DEBUG RENDERING: Processing chunk of size {len(chunk)}", flush=True)

                for local_idx, row in chunk.iterrows():

                    if stop_event and stop_event.is_set():

                        break

                    print(f"DEBUG RENDERING: Processing row {start_idx + processed + 1}", flush=True)

                    process_row(row.to_dict(), start_idx + processed)

                    processed += 1

                    if progress_callback:

                        progress_callback(total_processed + processed, total_to_process + (total_processed if total_processed>0 else 0))

                    if processed % 10 == 0:

                        with open(checkpoint_file, 'w') as f:

                            f.write(str(start_idx + processed))

                if stop_event and stop_event.is_set():

                    break



        if not (stop_event and stop_event.is_set()):

            if checkpoint_file.exists():

                checkpoint_file.unlink()

        total_processed += processed



    # Final: creare arhivă ZIP pentru tot output-ul

    zip_path_str = None

    if not (stop_event and stop_event.is_set()) and (zip_gen or send_mode == 'all_zip') and all_generated_files:

        zip_path = output_folder.parent / f"{output_folder.name}.zip"

        try:

            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:

                for file_path in all_generated_files:

                    if os.path.exists(file_path):

                        arcname = os.path.relpath(file_path, output_folder.parent)

                        zipf.write(file_path, arcname)

            zip_path_str = str(zip_path)

            if log_queue:

                log_queue.put(f"📦 Arhivă ZIP creată: {zip_path}")

        except Exception as e:

            if log_queue:

                log_queue.put(f"❌ Eroare la crearea arhivei: {e}")



    # Trimitere email arhivă totală (all_zip)

    if not (stop_event and stop_event.is_set()) and send_mode == 'all_zip' and zip_path_str and email_config and email_config.get('enabled') and email_config.get('to'):

        try:

            recipient_email = email_config['to']

            if not recipient_email or '@' not in str(recipient_email):

                if log_queue:

                    log_queue.put(f"⚠Adresă email destinatar invalidă: '{recipient_email}' - se omite trimiterea arhivei.")

            else:

                email_cfg = {

                    'smtp_server': email_config['smtp_server'],

                    'smtp_port': email_config['smtp_port'],

                    'username': email_config['username'],

                    'password': email_config['password'],

                    'from': email_config['from'],

                    'to': recipient_email,

                    'subject': email_config['subject'],

                    'body': email_config['body'],

                    'no_auth': email_config.get('no_auth', False)

                }

                if log_queue:

                    log_queue.put(f"📧 Trimite arhivă generală către {recipient_email} cu subiectul: {email_config['subject']}")

                send_email_with_attachments(email_cfg, [zip_path_str], log_queue, None)

        except Exception as e:

            error_msg = f"Eroare la trimiterea emailului cu arhiva generală: {e}"

            logger.error(error_msg)

            if log_queue:

                log_queue.put(f"❌ {error_msg}")



    # Merge documents (creare master doc)

    if not (stop_event and stop_event.is_set()) and merge_gen and HAS_MERGE and all_generated_files:

        try:

            master_doc = Document()

            docx_files = [f for f in all_generated_files if f.suffix == '.docx']

            for i, doc_path in enumerate(docx_files):

                sub_doc = Document(doc_path)

                for element in sub_doc.element.body:

                    master_doc.element.body.append(element)

                if i < len(docx_files) - 1:

                    master_doc.add_page_break()

            master_name = f"_MASTER_MERGED_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"

            master_path = output_folder / master_name

            master_doc.save(master_path)

            if log_queue:

                log_queue.put(f"🖨Document master creat: {master_path}")

        except Exception as e:

            if log_queue:

                log_queue.put(f"❌ Eroare la merge: {e}")



    if log_queue and not (stop_event and stop_event.is_set()):

        log_queue.put(f"Toate documentele au fost generate în: {output_folder}")

        if zip_path_str:

            log_queue.put(f"Arhivă ZIP creată: {zip_path_str}")

    logger.info(f"Generare finalizată. {total_processed} rânduri procesate.")