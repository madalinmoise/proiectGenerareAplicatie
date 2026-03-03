# app.py

import tkinter as tk

from typing import Set, List, Dict, Any, Optional

from tkinter import ttk, filedialog, messagebox, scrolledtext

import tkinter.simpledialog

import threading

import queue

import time

from pathlib import Path

from collections import deque

from docx import Document

from docxtpl import DocxTemplate

import traceback

import json

import os

import importlib.util

import re

import shutil

import webbrowser

import sys

import sqlite3

import subprocess

import zipfile

import web_api

import pandas as pd

import matplotlib.pyplot as plt

from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg

from datetime import datetime

import requests

import logging

from collections import Counter

import multiprocessing as mp

import pickle

import io

import gc

import hashlib

import customtkinter as ctk

from ui_constants import ThemeManager, PRIMARY_COLOR

from toast import show_toast



# Importuri proprii

from config import config

from audit import audit

from plugin_manager import plugin_manager

import version

import ui_utils

from ui_utils import add_tooltip, CreateToolTip, ScrollableFrame

from macro import MacroRecorder

from template_utils import (

    extract_all_placeholders_from_files,

    generate_excel_template,

    render_documents,

    render_document_from_template,

    extract_all_placeholders_async,

    extract_placeholders_from_file

)

from scanner_utils import (

    normalize_placeholder, scan_template_files,

    export_scan_report_html, export_scan_report_pdf

)

from email_utils import send_email_with_attachments

from ai_guide import LocalAIGuide

from multiprocessing_engine import MultiprocessingEngine

from performance_tracker import PerformanceTracker

from ui_batcher import UIUpdateBatcher



# Logger pentru aplicație

logger = logging.getLogger('HRAudit')



# Importuri opșionale

try:

    from tkinterdnd2 import TkinterDnD, DND_FILES

    HAS_DND = True

    Parent = TkinterDnD.Tk

except ImportError:

    HAS_DND = False

    Parent = tk.Tk



try:

    import pystray

    from PIL import Image, ImageDraw

    HAS_TRAY = True

except ImportError:

    HAS_TRAY = False



try:

    from flask import Flask

    HAS_FLASK = True

except ImportError:

    HAS_FLASK = False



try:

    from reportlab.lib.pagesizes import A4

    from reportlab.pdfgen import canvas

    HAS_REPORTLAB = True

except ImportError:

    HAS_REPORTLAB = False



try:

    import psutil

    HAS_PSUTIL = True

except ImportError:

    HAS_PSUTIL = False



try:

    from watchdog.observers import Observer

    from watchdog.events import FileSystemEventHandler

    HAS_WATCHDOG = True

except ImportError:

    HAS_WATCHDOG = False



try:

    import aiofiles

    import asyncio

    HAS_ASYNC = True

except ImportError:

    HAS_ASYNC = False



try:

    from odf.opendocument import load

    from odf.text import P

    HAS_ODF = True

except ImportError:

    HAS_ODF = False



try:

    from docx2pdf import convert as convert_to_pdf

    HAS_PDF = True

except ImportError:

    HAS_PDF = False



try:

    from docxcompose.composer import Composer

    HAS_MERGE = True

except ImportError:

    HAS_MERGE = False



# Importuri pentru tab-uri

from tabs.extract_tab import setup_extract_tab

from tabs.render_tab import setup_render_tab

from tabs.scan_tab import setup_scan_tab

from tabs.word_viewer_tab import setup_word_viewer_tab, refresh_word_file_list, highlight_placeholders_in_viewer

from tabs.excel_viewer_tab import setup_excel_viewer_tab, refresh_excel_viewer_content

from tabs.library_tab import setup_library_tab

from tabs.reports_tab import setup_reports_tab

from tabs.stats_tab import setup_stats_tab

from tabs.dashboard_tab import setup_dashboard_tab

from tabs.settings_tab import setup_settings_tab

from tabs.audit_tab import setup_audit_tab

from tabs.ai_guide_tab import setup_ai_guide_tab

from tabs.manual_tab import setup_manual_tab

from tabs.email_tab import setup_email_tab

# split_view_tab fuzionat în excel_viewer_tab



# Variabilă globală pentru script-uri (folosită în template_utils)

app_scripts = []



# ============================================================================

# PERFORMANCE OPTIMIZATION: Compiled Regex Patterns

# ============================================================================

TAG_PATTERN = re.compile(r'\{\{(.*?)\}\}')

WHITESPACE_PATTERN = re.compile(r'^\s*$')



# ============================================================================

# PERFORMANCE OPTIMIZATION: Template Tag Cache

# ============================================================================

class TemplateCache:

    """Cache for template analysis to avoid redundant parsing"""

    def __init__(self):

        self._cache: Dict[str, Set[str]] = {}

        self._lock = threading.Lock()

        self.quick_frame = None  # va fi setat în ai_guide_tab.py

        self.lbl_loading_excel = None  # va fi setat în excel_viewer_tab.py



    def get_tags(self, template_path: str, content: bytes = None) -> Set[str]:

        with self._lock:

            if template_path and content is None:

                try:

                    mtime = os.path.getmtime(template_path)

                    cache_key = f"{template_path}_{mtime}"

                    if cache_key in self._cache:

                        return self._cache[cache_key].copy()

                except:

                    pass



            tags = self._extract_tags_from_content(content or self._read_file(template_path))



            if template_path:

                try:

                    mtime = os.path.getmtime(template_path)

                    cache_key = f"{template_path}_{mtime}"

                    self._cache[cache_key] = tags

                except:

                    pass



            return tags.copy()



    def _read_file(self, path: str) -> bytes:

        with open(path, "rb") as f:

            return f.read()



    def _extract_tags_from_content(self, content: bytes) -> Set[str]:

        try:

            mem_file = io.BytesIO(content)

            doc = Document(mem_file)

            text_parts = []

            for p in doc.paragraphs:

                text_parts.append(p.text)

            for table in doc.tables:

                for row in table.rows:

                    for cell in row.cells:

                        text_parts.append(cell.text)

            full_text = ' '.join(text_parts)

            tags = {tag.strip() for tag in TAG_PATTERN.findall(full_text)}

            return tags

        except Exception as e:

            print(f"Error extracting tags: {e}")

            return set()



    def clear(self):

        with self._lock:

            self._cache.clear()



TEMPLATE_CACHE = TemplateCache()



# ============================================================================

# PERFORMANCE OPTIMIZATION: Encoding Constants

# ============================================================================

ENCODING_UTF8 = 'utf-8'

ENCODING_UTF8_SIG = 'utf-8-sig'



# --- Patch CTkScrollableFrame mouse wheel crash (event.widget can be a str) ---
def _patched_check_if_master_is_canvas(self, widget):
    try:
        if isinstance(widget, str):
            return False
        if widget == self._parent_canvas:
            return True
        elif widget.master is not None:
            return self.check_if_master_is_canvas(widget.master)
        else:
            return False
    except Exception:
        return False

try:
    import customtkinter.windows.widgets.ctk_scrollable_frame as _csf
    _csf.CTkScrollableFrame.check_if_master_is_canvas = _patched_check_if_master_is_canvas
except Exception:
    pass
# --- End patch ---

class SidebarNotebook(ctk.CTkFrame):

    def add(self, name):

        if name in self.tabs:

            return



        frame = ctk.CTkFrame(self.main_area, corner_radius=0, fg_color="transparent")

        self.tabs[name] = frame



        btn = ctk.CTkButton(self.sidebar_container, text=name, anchor="w",

                            fg_color="transparent", text_color=("gray10", "gray90"),

                            hover_color=("gray70", "gray30"), font=("Inter", 13),

                            command=lambda n=name: self.set(n))

        btn.pack(fill="x", pady=2, padx=5)

        self.buttons[name] = btn



    def tab(self, name):

        return self.tabs[name]



    def set(self, name):

        if self.current_tab == name:

            return

        for t_name, frame in self.tabs.items():

            frame.grid_forget()

            self.buttons[t_name].configure(fg_color="transparent")



        self.tabs[name].grid(row=0, column=0, sticky="nsew")

        self.buttons[name].configure(fg_color=("#4CAF50", "#2E7D32")) # Green accent for active tab

        self.current_tab = name



    def get(self):

        return self.current_tab





class PlaceholderApp(ctk.CTk):

    def show_disclaimer(self):

        """Fereastra de start cu termenii legali"""

        disc_win = tk.Toplevel(self)

        disc_win.title("TERMENI ȘI CONDIțII DE UTILIZARE")

        disc_win.geometry("750x700")

        disc_win.configure(bg="white")

        disc_win.grab_set()

        disc_win.attributes('-topmost', True)



        frame_logo = tk.Frame(disc_win, bg="white")

        frame_logo.pack(pady=20)

        self.draw_internal_logo(frame_logo)

        tk.Label(disc_win, text="APLICAțIE DE AUTOMATIZARE DOCUMENTE",

                 font=("Helvetica", 16, "bold"), bg="white", fg="#333").pack(pady=(10, 5))



        tk.Label(disc_win, text="Universitatea Nașională de Știinșă și Tehnologie POLITEHNICA București",

                 font=("Helvetica", 10), bg="white", fg="#555").pack()

        text_frame = tk.Frame(disc_win, bd=2, relief="groove")

        text_frame.pack(fill="both", expand=True, padx=30, pady=20)



        t = scrolledtext.ScrolledText(text_frame, height=10, font=("Arial", 11), wrap="word")

        t.pack(fill="both", expand=True)



        mesaj_legal = """

        NOTĂ DE COPYRIGHT ȘI EXONERARE DE RĂSPUNDERE



        1. DREPTURI DE AUTOR

        Această aplicație informatică este proprietatea intelectuală a S.l. dr. ing. Mădălin Vasile Moise.

        Toate drepturile asupra codului sursă, designului și logicii de funcționare îi aparșin autorului.



        2. LIMITAREA RĂSPUNDERII

        Această aplicație este oferită "ca atare" (as is), fără nicio garanție, explicită sau implicită.



        Autorul NU este răspunzător pentru:

        - Eventuala funcționalitate eronată a aplicației.

        - Erori în documentele generate.

        - Pierderi de date sau alte prejudicii cauzate de utilizarea software-ului.



        3. OBLIGAțIILE UTILIZATORULUI

        - Utilizatorul folosește această aplicație PE PROPRIA RĂSPUNDERE.

        - Toate informațiile și documentele generate TREBUIE REVERIFICATE MANUAL de către utilizator înainte de a fi semnate, transmise sau utilizate oficial.

        - Utilizatorul este singurul responsabil pentru validitatea juridică a conținutului documentelor rezultate.



        Prin apăsarea butonului de mai jos, confirmați că ați citit, ați înșeles și sunteși de acord cu acești termeni.

        """

        t.insert(tk.END, mesaj_legal)

        t.configure(state='disabled')

        btn_accept = tk.Button(disc_win, text="AM LUAT LA CUNOȘTINțĂ ȘI SUNT DE ACORD",

                               bg="#2E7D32", fg="white", font=("Arial", 11, "bold"), height=2,

                               command=lambda: self.start_app(disc_win))

        btn_accept.pack(fill="x", side="bottom", padx=30, pady=20)

        disc_win.protocol("WM_DELETE_WINDOW", self.quit_app)



    def start_app(self, window):

        if window:

            window.destroy()

        self.deiconify()

        self.log("Aplicație pornită. Realizat de Sef Lucrari Moise Madalin Vasile. Utilizați pe propria răspundere.")



    def quit_app(self):

        self.save_config()

        self.quit()



    def draw_internal_logo(self, parent):

        """Desenează o siglă vectorială"""

        canvas = tk.Canvas(parent, width=70, height=70, bg=parent.cget("bg"), highlightthickness=0)

        canvas.pack(side="left", padx=5)

        canvas.create_polygon(35, 5, 65, 15, 65, 45, 35, 65, 5, 45, 5, 15, fill="#002B7F", outline="white", width=2)

        canvas.create_text(35, 25, text="DOC", fill="white", font=("Arial", 9, "bold"))

        canvas.create_text(35, 40, text="GEN", fill="white", font=("Arial", 9, "bold"))



    # --------------------------------------------------------

    # Metode pentru toolbar și meniu

    # --------------------------------------------------------

    def create_toolbar(self):

        toolbar = ctk.CTkFrame(self, height=40, corner_radius=0)

        toolbar.pack(side='top', fill='x', padx=0, pady=0)



        btn_add = ctk.CTkButton(toolbar, text="Adaugă fișiere", width=120, command=self.add_template_files_menu)

        btn_add.pack(side='left', padx=5, pady=5)

        add_tooltip(btn_add, "Adaugă fișiere șablon (Ctrl+O)")



        btn_save = ctk.CTkButton(toolbar, text="Salvează config", width=120, command=self.save_config)

        btn_save.pack(side='left', padx=5, pady=5)

        add_tooltip(btn_save, "Salvează configurașia (Ctrl+S)")



        btn_render = ctk.CTkButton(toolbar, text="Generează", width=120, fg_color="#2e7d32", hover_color="#1b5e20", command=self.start_render)

        btn_render.pack(side='left', padx=5, pady=5)

        self.render_btn = btn_render

        add_tooltip(btn_render, "Începe generarea documentelor (Ctrl+R)")



        btn_extract = ctk.CTkButton(toolbar, text="Extrage", width=120, command=self.start_extract)

        btn_extract.pack(side='left', padx=5, pady=5)

        add_tooltip(btn_extract, "Extrage placeholder-uri (Ctrl+E)")



        btn_scan = ctk.CTkButton(toolbar, text="Scanează", width=120, command=self.start_scan_interactive)

        btn_scan.pack(side='left', padx=5, pady=5)

        add_tooltip(btn_scan, "Scanează șabloane (Ctrl+F)")



        btn_audit = ctk.CTkButton(toolbar, text="Audit", width=120, command=lambda: self.notebook.set("Audit Log"))

        btn_audit.pack(side='left', padx=5, pady=5)

        add_tooltip(btn_audit, "Vezi audit log (Ctrl+L)")



        # --- Font size controls ---
        font_ctrl_frame = ctk.CTkFrame(toolbar, fg_color="transparent")
        font_ctrl_frame.pack(side='right', padx=6, pady=5)

        self._current_font_size = 12

        btn_font_minus = ctk.CTkButton(font_ctrl_frame, text="A-", width=36,
                                       font=("Segoe UI", 11, "bold"),
                                       fg_color="#455a64", hover_color="#37474f",
                                       command=lambda: self.change_font_size(-1))
        btn_font_minus.pack(side='left', padx=2)
        add_tooltip(btn_font_minus, "Micșorează textul")

        self._font_size_label = ctk.CTkLabel(font_ctrl_frame, text="12pt",
                                              width=36, font=("Segoe UI", 11, "bold"))
        self._font_size_label.pack(side='left', padx=2)

        btn_font_plus = ctk.CTkButton(font_ctrl_frame, text="A+", width=36,
                                      font=("Segoe UI", 11, "bold"),
                                      fg_color="#455a64", hover_color="#37474f",
                                      command=lambda: self.change_font_size(1))
        btn_font_plus.pack(side='left', padx=2)
        add_tooltip(btn_font_plus, "Mărește textul")
        # --- End font size controls ---

        btn_help = ctk.CTkButton(toolbar, text="Ajutor", width=100, fg_color="gray", hover_color="#555", command=self.show_help)
        btn_help.pack(side='right', padx=10, pady=5)
        add_tooltip(btn_help, "Documentașie (F1)")



    def change_font_size(self, delta):
        """Modifică dinamic dimensiunea fontului pentru toată interfața."""
        new_size = self._current_font_size + delta
        if not (8 <= new_size <= 20):
            return
        self._current_font_size = new_size
        self._font_size_label.configure(text=f"{new_size}pt")

        from tkinter import font as tkfont
        # 1. Actualizează fontul global Tkinter (afectează toate widget-urile tk)
        for font_name in ("TkDefaultFont", "TkTextFont", "TkMenuFont", "TkHeadingFont"):
            try:
                f = tkfont.nametofont(font_name)
                f.configure(size=new_size)
            except Exception:
                pass
        try:
            tkfont.nametofont("TkFixedFont").configure(size=max(8, new_size - 1))
        except Exception:
            pass

        # 2. Actualizează ttk.Style (afectează ttk.Button, ttk.Label, ttk.Entry, Treeview etc.)
        from tkinter import ttk as _ttk
        _style = _ttk.Style()
        _F = ("Segoe UI", new_size)
        _FB = ("Segoe UI", new_size, "bold")
        for sty in (".", "TLabel", "TEntry", "TCheckbutton", "TRadiobutton", "TCombobox"):
            try:
                _style.configure(sty, font=_F)
            except Exception:
                pass
        for sty in ("TButton", "TLabelframe.Label", "Treeview.Heading"):
            try:
                _style.configure(sty, font=_FB)
            except Exception:
                pass
        try:
            row_h = max(20, new_size + 14)
            _style.configure("Treeview", font=_F, rowheight=row_h)
        except Exception:
            pass

        # 3. Actualizează widget-urile cu fonturi hardcodate
        for widget in [self.log_area, self.extract_placeholder_listbox,
                        self.scan_files_listbox, self.render_files_listbox,
                        self.extract_files_listbox]:
            try:
                if widget:
                    widget.configure(font=("Segoe UI", new_size))
            except Exception:
                pass

        # Forțează re-render
        self.update_idletasks()


    def create_menu(self):

        menubar = tk.Menu(self)

        self.configure(menu=menubar)



        file_menu = tk.Menu(menubar, tearoff=0)

        menubar.add_cascade(label="Fișier", menu=file_menu)

        file_menu.add_command(label="Salvează configurașia", command=self.save_config, accelerator="Ctrl+S")

        file_menu.add_command(label="Încarcă configurașia", command=self.load_config)

        file_menu.add_separator()

        file_menu.add_command(label="Salvează proiect ca...", command=self.save_project)

        file_menu.add_command(label="Deschide proiect...", command=self.load_project)

        file_menu.add_separator()

        file_menu.add_command(label="Adaugă fișiere șablon", command=self.add_template_files_menu, accelerator="Ctrl+O")

        file_menu.add_command(label="Ieșire", command=self.quit_app)



        recent_menu = tk.Menu(menubar, tearoff=0)

        menubar.add_cascade(label="Recente", menu=recent_menu)

        self.populate_recent_menu(recent_menu)



        tools_menu = tk.Menu(menubar, tearoff=0)

        menubar.add_cascade(label="Unelte", menu=tools_menu)

        tools_menu.add_command(label="Asistent creare șabloane", command=self.open_template_wizard)

        tools_menu.add_command(label="Istoric corecții", command=self.show_history_dialog)

        tools_menu.add_command(label="Plugin marketplace", command=self.open_plugin_marketplace)



        stats_menu = tk.Menu(menubar, tearoff=0)

        menubar.add_cascade(label="Statistici", menu=stats_menu)

        stats_menu.add_command(label="Actualizează statistici", command=self.update_stats)

        stats_menu.add_command(label="Exportă audit CSV", command=self.export_audit_csv)



        help_menu = tk.Menu(menubar, tearoff=0)

        menubar.add_cascade(label="Ajutor", menu=help_menu)

        help_menu.add_command(label="Despre", command=self.show_about)

        help_menu.add_command(label="Verifică actualizări", command=self.check_for_updates_now)

        help_menu.add_command(label="Documentașie", command=self.show_help, accelerator="F1")

        help_menu.add_command(label="Plugin API", command=lambda: webbrowser.open("docs/plugins.md"))



    def populate_recent_menu(self, menu):

        recent = config.get('recent_templates', [])

        menu.delete(0, tk.END)

        if not recent:

            menu.add_command(label="(niciun fișier recent)", state='disabled')

        else:

            for f in recent[:10]:

                menu.add_command(label=Path(f).name, command=lambda f=f: self.add_template_files([f]))



    def create_status_bar(self):

        self.status_frame = ctk.CTkFrame(self, height=30, corner_radius=0)

        self.status_frame.pack(side='bottom', fill='x')



        self.status_label = ctk.CTkLabel(self.status_frame, text="Pregătit", anchor='w')

        self.status_label.pack(side='left', fill='x', expand=True, padx=10)



        # Tema Switch (Nou în UI Upgrade)

        self.theme_btn = ctk.CTkButton(self.status_frame, text="Temă", width=80,

                                       command=self.toggle_app_theme)

        self.theme_btn.pack(side='right', padx=10, pady=5)



        self.progress_bar = ctk.CTkProgressBar(self.status_frame, width=200)

        self.progress_bar.pack(side='right', padx=10, pady=5)

        self.progress_bar.set(0)



        self.stop_button = ctk.CTkButton(self.status_frame, text="Oprește", width=80, fg_color="#c62828", hover_color="#b71c1c", command=self.stop_operation, state='disabled')

        self.stop_button.pack(side='right', padx=10, pady=5)



    def stop_button_toggle(self, state):

        self.stop_button.configure(state=state)



    def toggle_app_theme(self):

        new_mode = ThemeManager.toggle_theme()

        show_toast(self, f"Ați activat modul {new_mode}!")



    # --------------------------------------------------------

    # Metode ajutătoare (log, progres, etc.)

    # --------------------------------------------------------

    def log(self, msg):

        try:

            print(f"DEBUG LOG: {msg}", flush=True)

        except UnicodeEncodeError:

            # Fallback for consoles with limited encoding

            safe_msg = str(msg).encode('ascii', 'ignore').decode('ascii')

            print(f"DEBUG LOG (safe): {safe_msg}", flush=True)



        if hasattr(self, 'log_area') and self.log_area:

            self.log_area.insert(tk.END, msg + "\n")

            self.log_area.see(tk.END)

            self.update_idletasks()

        self.macro_recorder.record(msg)



    def log_hyperlink(self, label, filepath, tag):

        """Adaugă un link click-abil în log"""

        unique_tag = f"LINK:{filepath}|{tag}"

        def update():

            self.log_area.configure(state='normal')

            self.log_area.insert(tk.END, label, ("hyper", unique_tag))

            self.log_area.insert(tk.END, "\n")

            self.log_area.see(tk.END)

            self.log_area.configure(state='disabled')

        self.ui_batcher.add_update(update)



    def _on_log_click(self, event):

        try:

            index = self.log_area.index(f"@{event.x},{event.y}")

            tags = self.log_area.tag_names(index)

            for tag in tags:

                if tag.startswith("LINK:"):

                    _, payload = tag.split(":", 1)

                    if "|" in payload:

                        fpath, term = payload.split("|", 1)

                        if os.path.exists(fpath):

                            self.load_word_preview_highlight(fpath, term)

                        else:

                            messagebox.showwarning("Lipsă Fișier", f"Fișierul nu mai există:\n{fpath}")

                    return

        except Exception as e:

            print(f"Log click error: {e}")



    def load_word_preview_highlight(self, path, highlight_text=""):

        """Încarcă textul din Word și face highlight la termen (în tab-ul Word Viewer)"""

        self.notebook.set("Vizualizare documente")

        self.word_current_file = path

        self.word_search_var.set(highlight_text)

        self.word_viewer_load_file(path, highlight_placeholder=highlight_text)



    def poll_log_queue(self):

        while not self.log_queue.empty():

            msg = self.log_queue.get_nowait()

            if isinstance(msg, tuple):

                cmd = msg[0]

                if cmd == 'PROGRESS':

                    self.set_progress(msg[1], msg[2])

                elif cmd == 'GLOBAL_EXCEL_LOADED':

                    self.log(msg[1])

                    # Update all column combos

                    combos_to_update = []

                    if hasattr(self, 'email_column_combo'): combos_to_update.append(self.email_column_combo)

                    if hasattr(self, 'custom_x_combo'): combos_to_update.append(self.custom_x_combo)

                    if hasattr(self, 'custom_y_combo'): combos_to_update.append(self.custom_y_combo)

                    if hasattr(self, 'folder_column_combo'): combos_to_update.append(self.folder_column_combo)



                    if self.global_excel_df is not None:

                        cols = ["(Niciunul)"] + list(self.global_excel_df.columns)

                        for combo in combos_to_update:

                            combo['values'] = cols



                        # Special refresh for merged Excel Viewer

                        refresh_excel_viewer_content(self)

                        self.refresh_dashboard()

                        self.notebook.set("Dashboard")



                elif cmd == 'EXCEL_ERROR':

                    error_text = msg[1] if len(msg) > 1 else str(msg)

                    self.log(error_text)

                    messagebox.showerror("Eroare Excel", error_text)

                    if hasattr(self, 'lbl_loading_excel') and self.lbl_loading_excel:

                        self.lbl_loading_excel.configure(text="")

                elif cmd == 'FINISH':

                    self.is_generating = False

                    self.render_btn.configure(state='normal')

                    self.batch_render_btn.configure(state='normal')

                    count = msg[1]

                    stats = msg[2]

                    self.log(f"✅ Proces complet! Succes: {stats['success']}, Erori: {stats['errors']}, Timp: {stats['duration']:.2f}s")

                    if stats.get('zip_path'):

                        if messagebox.askyesno("Arhivă ZIP", f"Arhiva a fost creată la:\n{stats['zip_path']}\nDoriși să o deschideși?"):

                            os.startfile(stats['zip_path'])

                    # Refresh word viewer list

                    refresh_word_file_list(self)

                elif cmd == 'ERROR_BOX':

                    messagebox.showerror("Eroare", msg[1])

                elif isinstance(msg, tuple) and msg[0] == 'SUCCESS':

                    self.success_count += 1

                    self.set_progress(self.progress_last_count, self.total_count)

                elif isinstance(msg, tuple) and msg[0] == 'ERROR':

                    self.error_count += 1

                    self.set_progress(self.progress_last_count, self.total_count)

                else:

                    self.log(str(msg))

            else:

                self.log(msg)

        self.after(100, self.poll_log_queue)



    def set_progress(self, current, total):

        if total > 0:

            self.progress_val = current / total

        else:

            self.progress_val = 0



        self.total_count = total



        if hasattr(self, 'progress_bar'):
            try:
                self.progress_bar.set(self.progress_val)
            except AttributeError:
                self.progress_bar['value'] = self.progress_val * 100



        now = time.time()

        if self.progress_start_time is None:

            self.progress_start_time = now

            self.progress_last_count = current

            self.progress_last_time = now

            remaining_str = "?"

        else:

            elapsed = now - self.progress_last_time

            if elapsed > 0:

                speed = (current - self.progress_last_count) / elapsed

                if speed > 0:

                    remaining = (total - current) / speed

                    remaining_str = time.strftime("%H:%M:%S", time.gmtime(remaining))

                else:

                    remaining_str = "?"

            else:

                remaining_str = "?"

            self.progress_last_count = current

            self.progress_last_time = now



        self.status_label['text'] = f"Progres: {current}/{total}  Rămas: {remaining_str}"

        self.progress_last_count = current



        # Update Dashboard Widgets (Feature 5)

        if hasattr(self, 'lbl_stat_total'):

            self.lbl_stat_total.configure(text=f"Total: {total}")

        if hasattr(self, 'lbl_stat_success'):

            self.lbl_stat_success.configure(text=f"Succes: {self.success_count}")

        if hasattr(self, 'lbl_stat_error'):

            self.lbl_stat_error.configure(text=f"Erori: {self.error_count}")

        if hasattr(self, 'progress_bar'):

            self.progress_bar['value'] = (current / total) * 100 if total > 0 else 0



        self.update_idletasks()



    def reset_progress(self):

        self.progress_bar['value'] = 0

        self.status_label['text'] = "Pregătit"

        self.stop_button.configure(state='disabled')

        self.stop_render_event.clear()

        self.progress_start_time = None

        self.progress_last_count = 0

        self.progress_last_time = None



    def notify_done(self, message="Operașiune finalizată!"):

        self.log(message)

        # S-a scos messagebox și winsound la cererea utilizatorului



    def apply_theme(self):

        style = ttk.Style()

        style.configure('White.TFrame', background='white')

        theme = self.theme.get()

        if theme == "dark":

            style.theme_use('clam')

            bg = '#2d2d2d'

            fg = '#ffffff'

            entry_bg = '#3c3c3c'

            select_bg = '#4a4a4a'

            _F = ('Segoe UI', 12)
            _FB = ('Segoe UI', 12, 'bold')
            style.configure('.', background=bg, foreground=fg, fieldbackground=entry_bg, font=_F)

            style.configure('TLabel', background=bg, foreground=fg, font=_F)

            style.configure('TButton', background=entry_bg, foreground=fg, font=_FB, padding=6)

            style.map('TButton', background=[('active', select_bg)])

            style.configure('TFrame', background=bg)

            style.configure('TLabelframe', background=bg, foreground=fg)

            style.configure('TLabelframe.Label', background=bg, foreground=fg, font=_FB)

            style.configure('TNotebook', background=bg, foreground=fg)

            style.configure('TNotebook.Tab', background=entry_bg, foreground=fg, font=_FB, padding=[10, 5])

            style.configure('Treeview', font=_F, rowheight=28)

            style.configure('Treeview.Heading', font=_FB)

            style.configure('TCheckbutton', font=_F)

            style.configure('TRadiobutton', font=_F)

            style.map('TNotebook.Tab', background=[('selected', select_bg)])

            self.configure(bg=bg)

            for widget in [self.log_area, self.extract_placeholder_listbox, self.scan_files_listbox,

                           self.render_files_listbox, self.extract_files_listbox, self.word_text]:

                if widget:

                    try:

                        widget.configure(bg=entry_bg, fg=fg, insertbackground='white')

                    except:

                        pass

        else:

            style.theme_use('vista')

            _F = ('Segoe UI', 12)
            _FB = ('Segoe UI', 12, 'bold')
            style.configure('.', background='#f0f0f0', foreground='#000000', font=_F)

            style.configure('TLabel', font=_F)

            style.configure('TButton', font=_FB, padding=6)

            style.configure('TLabelframe.Label', font=_FB)

            style.configure('Treeview', font=_F, rowheight=28)

            style.configure('Treeview.Heading', font=_FB)

            style.configure('TCheckbutton', font=_F)

            style.configure('TRadiobutton', font=_F)

            self.configure(bg='#f0f0f0')

            for widget in [self.log_area, self.extract_placeholder_listbox, self.scan_files_listbox,

                           self.render_files_listbox, self.extract_files_listbox, self.word_text]:

                if widget:

                    try:

                        widget.configure(bg='white', fg='black', insertbackground='black')

                    except:

                        pass



    def apply_touch_mode(self):

        if self.touch_mode.get():

            default_font = ('Arial', 12)

            self.option_add('*Font', default_font)

            style = ttk.Style()

            style.configure('TButton', padding=10)

            style.configure('TLabel', padding=5)

            style.configure('TEntry', padding=5)

            style.configure('TCombobox', padding=5)

            style.configure('TCheckbutton', padding=5)

            style.configure('TRadiobutton', padding=5)

        else:

            self.option_add('*Font', ('Arial', 9))

            style = ttk.Style()

            style.configure('TButton', padding=2)

            style.configure('TLabel', padding=2)

            style.configure('TEntry', padding=2)

            style.configure('TCombobox', padding=2)

            style.configure('TCheckbutton', padding=2)

            style.configure('TRadiobutton', padding=2)



    def refresh_all_file_listboxes(self):

        for name in ['extract', 'render', 'scan']:

            listbox = getattr(self, f'{name}_files_listbox', None)

            if listbox:

                listbox.delete(0, tk.END)

                for f in self.template_files:

                    listbox.insert(tk.END, Path(f).name)



    def add_template_files(self, files):

        added = False

        for f in files:

            if f not in self.template_files:

                self.template_files.append(f)

                added = True

                self.macro_recorder.record(f"Adăugat fișier: {f}", f"app.add_template_files(['{f}'])")

        if added:

            self.refresh_all_file_listboxes()

            recent = config.get('recent_templates', [])

            for f in files:

                if f in recent:

                    recent.remove(f)

                recent.insert(0, f)

            config.set('recent_templates', recent[:20])

            self.log(f"Adăugate {len(files)} fișiere. Total: {len(self.template_files)}")

            audit.log(action='add_templates', details={'count': len(files)})
        if hasattr(self, 'engine'): self.engine.notify_observers('config_updated', {'template_files': self.template_files})



    def remove_selected_files(self, listbox):

        selected = listbox.curselection()

        if not selected:

            return

        selected_names = [listbox.get(i) for i in selected]

        new_list = []

        for f in self.template_files:

            if Path(f).name not in selected_names:

                new_list.append(f)

        self.template_files = new_list

        self.refresh_all_file_listboxes()

        self.log(f"Eliminate {len(selected)} fișiere. Rămân {len(self.template_files)}")

        audit.log(action='remove_templates', details={'count': len(selected)})

        self.macro_recorder.record(f"Eliminate {len(selected)} fișiere")



    def setup_drop_target(self, widget, callback):

        if HAS_DND:

            try:

                widget.drop_target_register(DND_FILES)

                widget.dnd_bind('<<Drop>>', callback)

            except Exception as e:

                self.log(f"DnD indisponibil: {e}")



    def make_listbox_drag_drop(self, listbox):

        drag_start_index = None



        def on_drag_start(event):

            nonlocal drag_start_index

            index = listbox.nearest(event.y)

            if index >= 0:

                drag_start_index = index

                listbox.selection_clear(0, tk.END)

                listbox.selection_set(index)



        def on_drag_motion(event):

            nonlocal drag_start_index

            if drag_start_index is not None:

                index = listbox.nearest(event.y)

                if index >= 0 and index != drag_start_index:

                    # Mută elementul

                    item = listbox.get(drag_start_index)

                    listbox.delete(drag_start_index)

                    listbox.insert(index, item)

                    files = self.template_files.copy()

                    f = files.pop(drag_start_index)

                    files.insert(index, f)

                    self.template_files = files

                    self.refresh_all_file_listboxes()

                    drag_start_index = index



        def on_drag_end(event):

            nonlocal drag_start_index

            drag_start_index = None



        listbox.bind('<Button-1>', on_drag_start)

        listbox.bind('<B1-Motion>', on_drag_motion)

        listbox.bind('<ButtonRelease-1>', on_drag_end)



    def add_context_menu(self, listbox, add_func, remove_func):

        menu = tk.Menu(self, tearoff=0)

        menu.add_command(label="Adaugă fișiere", command=add_func)

        menu.add_command(label="Șterge selectate", command=lambda: remove_func(listbox))



        def show_menu(event):

            menu.post(event.x_root, event.y_root)



        listbox.bind('<Button-3>', show_menu)



    def bind_all_shortcuts(self):

        self.bind_all('<Control-o>', lambda e: self.add_template_files_menu())

        self.bind_all('<Control-s>', lambda e: self.save_config())

        self.bind_all('<Control-r>', lambda e: self.start_render())

        self.bind_all('<Control-e>', lambda e: self.start_extract())

        self.bind_all('<Control-f>', lambda e: self.notebook.select(self.frame_scan))

        self.bind_all('<Control-l>', lambda e: self.notebook.select(self.frame_audit))

        self.bind_all('<F1>', lambda e: self.show_help())



    def show_help(self):

        doc_path = Path('docs/plugins.md')

        if doc_path.exists():

            webbrowser.open(str(doc_path))

        else:

            webbrowser.open("https://github.com/your-repo/wiki")



    def add_template_files_menu(self):

        files = filedialog.askopenfilenames(filetypes=[("Template files", "*.docx *.odt *.txt *.html")])

        if files:

            self.add_template_files(files)



    def start_auto_save(self):

        self.save_config()

        self.after(5 * 60 * 1000, self.auto_save)



    def auto_save(self):

        self.save_config()

        self.log("Auto-save: configurație salvată.")

        self.after(5 * 60 * 1000, self.auto_save)



    def save_config(self):

        config.set('template_files', self.template_files)

        config.set('excel_output', self.excel_output_path.get())

        config.set('data_file', self.data_file_path.get())

        config.set('output_dir', self.output_dir_path.get())

        config.set('folder_column', self.folder_column.get())

        config.set('filename_pattern', self.filename_pattern.get())

        config.set('theme', self.theme.get())

        config.set('chunksize', self.chunksize.get())

        config.set('touch_mode', self.touch_mode.get())

        config.set('email_enabled', self.email_config['enabled'].get())

        config.set('smtp_server', self.email_config['smtp_server'].get())

        config.set('smtp_port', self.email_config['smtp_port'].get())

        config.set('email_username', self.email_config['username'].get())

        config.set('email_password', self.email_config['password'].get())

        config.set('email_from', self.email_config['from'].get())

        config.set('email_to', self.email_config['to'].get())

        config.set('email_subject', self.email_config['subject'].get())

        config.set('email_body', self.email_config['body'].get())

        config.set('email_send_mode', self.email_send_mode.get())

        config.set('email_column', self.email_column.get())

        config.set('email_subject_pattern', self.email_subject_pattern.get())

        config.set('multiprocessing', self.multiprocessing_var.get())

        config.set('num_workers', self.num_workers_var.get())

        config.set('auto_recovery', self.auto_recovery_var.get())

        config.set('pdf_gen', self.pdf_var.get())

        config.set('zip_gen', self.zip_var.get())

        config.set('zip_per_row', self.zip_per_row_var.get())

        config.set('merge_gen', self.merge_var.get())

        config.set('audio_alert', self.audio_var.get())

        config.set('subfolder_col', self.subfolder_var.get())

        config.set('clean_data', self.clean_data_var.get())

        config.set('recent_files', self.recent_files)

        self.log("Configurație salvată.")



    def load_config(self):

        self.template_files = config.get('template_files', [])

        self.excel_output_path.set(config.get('excel_output', 'placeholders.xlsx'))

        self.data_file_path.set(config.get('data_file', ''))

        self.output_dir_path.set(config.get('output_dir', 'output'))

        self.folder_column.set(config.get('folder_column', ''))

        self.filename_pattern.set(config.get('filename_pattern', '{ID}_{template_name}.docx'))

        self.theme.set(config.get('theme', 'light'))

        self.chunksize.set(config.get('chunksize', 1000))

        self.touch_mode.set(config.get('touch_mode', False))

        self.email_config['enabled'].set(config.get('email_enabled', False))

        self.email_config['smtp_server'].set(config.get('smtp_server', ''))

        self.email_config['smtp_port'].set(config.get('smtp_port', 587))

        self.email_config['username'].set(config.get('email_username', ''))

        self.email_config['password'].set(config.get('email_password', ''))

        self.email_config['from'].set(config.get('email_from', ''))

        self.email_config['to'].set(config.get('email_to', ''))

        self.email_config['subject'].set(config.get('email_subject', 'Documente generate'))

        self.email_config['body'].set(config.get('email_body', 'Vă rugăm găsiși atașat documentele generate.'))

        self.email_send_mode.set(config.get('email_send_mode', 'individual'))

        self.email_column.set(config.get('email_column', ''))

        self.email_subject_pattern.set(config.get('email_subject_pattern', 'Documente generate pentru {ID}'))

        self.multiprocessing_var.set(config.get('multiprocessing', False))

        self.num_workers_var.set(config.get('num_workers', max(1, mp.cpu_count()-1)))

        self.auto_recovery_var.set(config.get('auto_recovery', False))

        self.pdf_var.set(config.get('pdf_gen', False))

        self.zip_var.set(config.get('zip_gen', False))

        self.zip_per_row_var.set(config.get('zip_per_row', False))

        self.merge_var.set(config.get('merge_gen', False))

        self.audio_var.set(config.get('audio_alert', True))

        self.subfolder_var.set(config.get('subfolder_col', '(Niciunul)'))

        self.clean_data_var.set(config.get('clean_data', True))

        self.recent_files = config.get('recent_files', [])

        self.refresh_all_file_listboxes()

        self.apply_theme()

        self.apply_touch_mode()

        self.log("Configurație încărcată.")



    def save_project(self):

        file_path = filedialog.asksaveasfilename(defaultextension=".hrproj", filetypes=[("HR Project", "*.hrproj")])

        if not file_path:

            return

        project = {

            'template_files': self.template_files,

            'excel_output': self.excel_output_path.get(),

            'data_file': self.data_file_path.get(),

            'output_dir': self.output_dir_path.get(),

            'folder_column': self.folder_column.get(),

            'filename_pattern': self.filename_pattern.get(),

            'theme': self.theme.get(),

            'chunksize': self.chunksize.get(),

            'touch_mode': self.touch_mode.get(),

            'email_enabled': self.email_config['enabled'].get(),

            'smtp_server': self.email_config['smtp_server'].get(),

            'smtp_port': self.email_config['smtp_port'].get(),

            'email_username': self.email_config['username'].get(),

            'email_password': self.email_config['password'].get(),

            'email_from': self.email_config['from'].get(),

            'email_to': self.email_config['to'].get(),

            'email_subject': self.email_config['subject'].get(),

            'email_body': self.email_config['body'].get(),

            'email_send_mode': self.email_send_mode.get(),

            'email_column': self.email_column.get(),

            'email_subject_pattern': self.email_subject_pattern.get(),

            'multiprocessing': self.multiprocessing_var.get(),

            'num_workers': self.num_workers_var.get(),

            'auto_recovery': self.auto_recovery_var.get(),

            'pdf_gen': self.pdf_var.get(),

            'zip_gen': self.zip_var.get(),

            'merge_gen': self.merge_var.get(),

            'audio_alert': self.audio_var.get(),

            'subfolder_col': self.subfolder_var.get(),

            'clean_data': self.clean_data_var.get(),

        }

        if hasattr(self, 'sheet_var'):

            project['sheet_name'] = self.sheet_var.get()

        with open(file_path, 'w', encoding='utf-8') as f:

            json.dump(project, f, indent=2)

        self.log_queue.put(f"Proiect salvat: {file_path}")



    def load_project(self):

        file_path = filedialog.askopenfilename(filetypes=[("HR Project", "*.hrproj")])

        if not file_path:

            return

        with open(file_path, 'r', encoding='utf-8') as f:

            project = json.load(f)

        self.template_files = project.get('template_files', [])

        self.excel_output_path.set(project.get('excel_output', 'placeholders.xlsx'))

        self.data_file_path.set(project.get('data_file', ''))

        self.output_dir_path.set(project.get('output_dir', 'output'))

        self.folder_column.set(project.get('folder_column', ''))

        self.filename_pattern.set(project.get('filename_pattern', '{ID}_{template_name}.docx'))

        self.theme.set(project.get('theme', 'light'))

        self.chunksize.set(project.get('chunksize', 1000))

        self.touch_mode.set(project.get('touch_mode', False))

        self.email_config['enabled'].set(project.get('email_enabled', False))

        self.email_config['smtp_server'].set(project.get('smtp_server', ''))

        self.email_config['smtp_port'].set(project.get('smtp_port', 587))

        self.email_config['username'].set(project.get('email_username', ''))

        self.email_config['password'].set(project.get('email_password', ''))

        self.email_config['from'].set(project.get('email_from', ''))

        self.email_config['to'].set(project.get('email_to', ''))

        self.email_config['subject'].set(project.get('email_subject', 'Documente generate'))

        self.email_config['body'].set(project.get('email_body', 'Vă rugăm găsiși atașat documentele generate.'))

        self.email_send_mode.set(project.get('email_send_mode', 'individual'))

        self.email_column.set(project.get('email_column', ''))

        self.email_subject_pattern.set(project.get('email_subject_pattern', 'Documente generate pentru {ID}'))

        self.multiprocessing_var.set(project.get('multiprocessing', False))

        self.num_workers_var.set(project.get('num_workers', max(1, mp.cpu_count()-1)))

        self.auto_recovery_var.set(project.get('auto_recovery', False))

        self.pdf_var.set(project.get('pdf_gen', False))

        self.zip_var.set(project.get('zip_gen', False))

        self.merge_var.set(project.get('merge_gen', False))

        self.audio_var.set(project.get('audio_alert', True))

        self.subfolder_var.set(project.get('subfolder_col', '(Niciunul)'))

        self.clean_data_var.set(project.get('clean_data', True))

        if hasattr(self, 'sheet_var') and 'sheet_name' in project:

            self.sheet_var.set(project['sheet_name'])

        self.refresh_all_file_listboxes()

        self.apply_theme()

        self.apply_touch_mode()

        self.log_queue.put(f"Proiect încărcat: {file_path}")



    # ========================================================================

    # AI GUIDE METHODS

    # ========================================================================

    def send_guide_message(self, predefined_msg=None):

        """Trimite mesaj către ghidul local"""

        if predefined_msg:

            user_message = predefined_msg

        else:

            user_message = self.guide_input.get().strip()

            if not user_message:

                return

            self.guide_input.delete(0, tk.END)



        self.add_guide_message("user", f"\nTU: {user_message}\n")

        response = self.local_guide.get_response(user_message)

        self.add_guide_message("guide", f"\nGHID: {response}\n")



        if self.local_guide.tutorial_active:

            step = self.local_guide.tutorial_steps[self.local_guide.current_step]

            self.tutorial_status_label.configure(

                text=f"TUTORIAL ACTIV: {step['title']} ({self.local_guide.current_step + 1}/{len(self.local_guide.tutorial_steps)})"

            )

            self.tutorial_status_frame.pack(fill='x', pady=(0, 10), after=self.quick_frame)

        else:

            self.tutorial_status_frame.pack_forget()



        self.guide_text.see(tk.END)



    def add_guide_message(self, tag, message):

        """Adaugă mesaj în chat"""

        self.guide_text.configure(state='normal')

        self.guide_text.insert(tk.END, message, tag)

        self.guide_text.configure(state='disabled')

        self.guide_text.see(tk.END)



    def clear_guide_conversation(self):

        if messagebox.askyesno("Confirmare", "Ștergi întreaga conversașie cu ghidul?\nAcest lucru va reseta și tutorial-ul dacă este activ."):

            self.local_guide.conversation_history = []

            self.local_guide.tutorial_active = False

            self.local_guide.current_step = 0

            self.guide_text.configure(state='normal')

            self.guide_text.delete(1.0, tk.END)

            self.guide_text.configure(state='disabled')

            self.tutorial_status_frame.pack_forget()

            self.add_guide_message("system", "Conversașie ștearsă. Poși începe din nou.")



    # ========================================================================

    # DASHBOARD METHODS

    # ========================================================================

    def refresh_dashboard(self):

        """Actualizează dashboard-ul cu datele curente din Excel"""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            return



        rows, cols = data.shape

        empty_count = data.isna().sum().sum() + (data == "").sum().sum()



        # Actualizăm label-urile dacă există

        if hasattr(self, 'lbl_stat_rows'):

            self.lbl_stat_rows.configure(text=f"Total Rânduri: {rows}")

        if hasattr(self, 'lbl_stat_cols'):

            self.lbl_stat_cols.configure(text=f"Coloane: {cols}")

        if hasattr(self, 'lbl_stat_empty'):

            self.lbl_stat_empty.configure(text=f"Celule Goale: {empty_count}")



        if hasattr(self, 'stats_text'):

            self.update_stats()



        # Ștergem canvas-ul vechi

        if not hasattr(self, 'dashboard_chart_frame'):

            return



        for widget in self.dashboard_chart_frame.winfo_children():

            widget.destroy()



        # Selectăm o coloană relevantă pentru grafic

        target_col = None

        # Priorităși pentru dashboard: post, departament, educație

        priorities = ['denumire_post', 'grad_profesional', 'Universitatea/departamentul', 'Required_Education_Level']

        for p in priorities:

            if p in data.columns:

                target_col = p

                break



        if not target_col:

            # Fallback la prima coloană de tip obiect (text)

            for c in data.columns:

                if data[c].dtype == 'object':

                    target_col = c

                    break



        fig, ax = plt.subplots(figsize=(10, 5))

        if target_col:

            counts = data[target_col].value_counts().head(10)

            counts.plot(kind='bar', ax=ax, color='#3498db', edgecolor='black')

            ax.set_title(f'Distribuție: {target_col}', fontsize=12, fontweight='bold')

            ax.set_ylabel('Număr persoane', fontsize=10)

            ax.tick_params(axis='x', rotation=45)

            plt.grid(axis='y', linestyle='--', alpha=0.7)

            plt.tight_layout()



            canvas = FigureCanvasTkAgg(fig, master=self.dashboard_chart_frame)

            canvas.draw()

            canvas.get_tk_widget().pack(fill='both', expand=True, padx=10, pady=10)

        else:

            ax.text(0.5, 0.5, "Nu există coloane date pentru grafic", ha='center', va='center')

            canvas = FigureCanvasTkAgg(fig, master=self.dashboard_chart_frame)

            canvas.draw()

            canvas.get_tk_widget().pack(fill='both', expand=True)



        plt.close(fig)



    # ========================================================================

    # FORMS METHODS

    # ========================================================================

    def toggle_forms_cols(self, state):

        for var in self.forms_vars.values():

            var.set(state)



    def genereaza_fisier_forms(self):

        """Generează un fișier Word pentru import în Microsoft Forms"""

        active_cols = [col for col, var in self.forms_vars.items() if var.get()]

        if not active_cols:

            messagebox.showwarning("Atenție", "Nu ai selectat nicio coloană!")

            return

        f = filedialog.asksaveasfilename(defaultextension=".docx",

                                         filetypes=[("Word Document", "*.docx")],

                                         initialfile="Import_Forms.docx")

        if not f:

            return

        try:

            doc = Document()

            doc.add_heading('Formular Date', 0)

            p = doc.add_paragraph("Acest fișier este optimizat pentru importul în Microsoft Forms (Format Simplu).")

            p.italic = True

            doc.add_paragraph("")

            for index, col in enumerate(active_cols, 1):

                p = doc.add_paragraph(f"{index}. {col}")

                p.bold = True

                doc.add_paragraph("")

                doc.add_paragraph("")

            doc.save(f)

            messagebox.showinfo("Succes", f"Fișier generat!\nÎncarcă-l în Microsoft Forms folosind 'Import Rapid'.")

            self.log(f"Generat Forms Import: {f}")

        except Exception as e:

            messagebox.showerror("Eroare", str(e))



    # ========================================================================

    # MANUAL METHODS

    # ========================================================================

    def export_manual_pdf(self):

        try:

            from reportlab.lib.pagesizes import A4

            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer

            from reportlab.lib.styles import getSampleStyleSheet

            from reportlab.lib.units import inch



            filename = filedialog.asksaveasfilename(defaultextension=".pdf",

                                                     filetypes=[("PDF files", "*.pdf")],

                                                     initialfile="Manual_Operator_V6.0.pdf")

            if filename:

                doc = SimpleDocTemplate(filename, pagesize=A4)

                story = []

                styles = getSampleStyleSheet()

                content = self.manual_text.get("1.0", "end-1c")

                for line in content.split('\n'):

                    if line.strip():

                        story.append(Paragraph(line, styles['Normal']))

                        story.append(Spacer(1, 0.1*inch))

                doc.build(story)

                messagebox.showinfo("Succes", f"Manual salvat ca PDF:\n{filename}")

        except ImportError:

            messagebox.showwarning("Lipsă Bibliotecă", "Pentru export PDF, instalași: pip install reportlab")

        except Exception as e:

            messagebox.showerror("Eroare", f"Nu s-a putut exporta PDF:\n{e}")



    def export_manual_txt(self):

        filename = filedialog.asksaveasfilename(defaultextension=".txt",

                                                 filetypes=[("Text files", "*.txt")],

                                                 initialfile="Manual_Operator_V6.0.txt")

        if filename:

            try:

                content = self.manual_text.get("1.0", "end-1c")

                with open(filename, 'w', encoding='utf-8') as f:

                    f.write(content)

                messagebox.showinfo("Succes", f"Manual salvat ca TXT:\n{filename}")

            except Exception as e:

                messagebox.showerror("Eroare", f"Nu s-a putut salva:\n{e}")



    def search_in_manual(self):

        search_term = tk.simpledialog.askstring("Căutare în Manual", "Introduceși textul de căutat:")

        if search_term:

            self.manual_text.tag_remove("search", "1.0", "end")

            start = "1.0"

            count = 0

            while True:

                pos = self.manual_text.search(search_term, start, stopindex="end", nocase=True)

                if not pos:

                    break

                end = f"{pos}+{len(search_term)}c"

                self.manual_text.tag_add("search", pos, end)

                count += 1

                start = end

            self.manual_text.tag_config("search", background="yellow", foreground="black")

            if count > 0:

                first = self.manual_text.search(search_term, "1.0", stopindex="end", nocase=True)

                if first:

                    self.manual_text.see(first)

                messagebox.showinfo("Căutare", f"Găsite {count} rezultate pentru '{search_term}'")

            else:

                messagebox.showinfo("Căutare", f"Nu s-au găsit rezultate pentru '{search_term}'")



    # ========================================================================

    # EXCEL VIEWER METHODS (extinse)

    # ========================================================================

    def incarca_previzualizare_excel_async(self):

        if not self.data_file_path.get():

            return

        path = self.data_file_path.get()

        if not os.path.exists(path):

            return

        self.is_loading_excel = True

        self.lbl_loading_excel.configure(text="Se încarcă datele...")

        thread = threading.Thread(target=self._load_excel_thread, args=(path,))

        thread.daemon = True

        thread.start()



    def _load_excel_thread(self, path):

        try:

            df = pd.read_excel(path, dtype=str).fillna('')

            df.columns = df.columns.str.strip()

            self.log_queue.put(('EXCEL_LOADED', df))

        except Exception as e:

            self.log_queue.put(('EXCEL_ERROR', str(e)))



    def _populate_excel_tree(self, df):

        self.lbl_loading_excel.configure(text="")

        self.full_df = df



        # Stergem widgeturile vechi

        if hasattr(self, 'excel_columns_inner'):

            for w in self.excel_columns_inner.winfo_children():

                w.destroy()



        self.column_vars = {}

        for col in df.columns:

            var = tk.BooleanVar(value=True)

            self.column_vars[col] = var

            if hasattr(self, 'excel_columns_inner'):

                cb = tk.Checkbutton(self.excel_columns_inner, text=col, variable=var,

                                    command=self.refresh_excel_view, anchor="w", bg="white")

                cb.pack(fill="x", padx=2)



        self.cols_for_subfolder = ["(Niciunul)"] + list(df.columns)

        self.refresh_excel_view()

        self.log(f"Vizualizare Excel: {len(df)} rânduri încărcate.")



    def refresh_excel_view(self):

        if self.full_df is None:

            return

        active_cols = [col for col, var in self.column_vars.items() if var.get()]

        if not active_cols:

            self.excel_tree.delete(*self.excel_tree.get_children())

            return

        term = self.search_var.get().strip().lower()

        filtered_df = self.full_df[active_cols].copy()

        if term:

            mask = filtered_df.apply(lambda row: row.astype(str).str.lower().str.contains(term, na=False).any(), axis=1)

            filtered_df = filtered_df[mask]

        self.excel_tree.delete(*self.excel_tree.get_children())

        self.excel_tree["columns"] = active_cols

        self.excel_tree["show"] = "headings"

        for col in active_cols:

            self.excel_tree.heading(col, text=col)

            self.excel_tree.column(col, width=120)

        limit = 1000

        for idx, row in filtered_df.iloc[:limit].iterrows():

            self.excel_tree.insert("", "end", values=list(row))

        if len(filtered_df) > limit:

            self.lbl_loading_excel.configure(text=f"Afișate primele {limit} din {len(filtered_df)} rânduri.")



    def perform_search_excel(self):

        self.refresh_excel_view()



    def clear_search_excel(self):

        self.search_var.set("")

        self.refresh_excel_view()



    def toggle_all_cols(self, state):

        for var in self.column_vars.values():

            var.set(state)

        self.refresh_excel_view()



    def export_excel_visible(self):

        if self.full_df is None:

            return

        active_cols = [col for col, var in self.column_vars.items() if var.get()]

        if not active_cols:

            messagebox.showwarning("Export", "Nicio coloană selectată!")

            return

        term = self.search_var.get().strip().lower()

        filtered_df = self.full_df[active_cols].copy()

        if term:

            mask = filtered_df.apply(lambda row: row.astype(str).str.lower().str.contains(term, na=False).any(), axis=1)

            filtered_df = filtered_df[mask]

        f = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel", "*.xlsx")])

        if f:

            try:

                filtered_df.to_excel(f, index=False)

                messagebox.showinfo("Export", "Datele filtrate au fost exportate cu succes.")

            except Exception as e:

                messagebox.showerror("Eroare Export", str(e))



    def valideaza_datele(self):

        if self.full_df is None:

            return

        active_cols = [col for col, var in self.column_vars.items() if var.get()]

        if not active_cols:

            self.log("⚠️ Nicio coloană activă pentru validare.")

            return

        df_val = self.full_df[active_cols].replace(r'^\s*$', float('nan'), regex=True)

        issues = 0

        for col in active_cols:

            empty_rows = df_val[df_val[col].isna()].index.tolist()

            if empty_rows:

                excel_rows = [r+2 for r in empty_rows]

                self.log(f"⚠️ Coloana [{col}]: Valori lipsă la rândurile {excel_rows[:10]}...")

                issues += len(empty_rows)

        if issues == 0:

            self.log("✅ Datele sunt valide.")

            messagebox.showinfo("Validare", "Datele sunt OK! Nu s-au găsit celule goale.")

        else:

            self.log(f"❌ S-au găsit {issues} probleme.")

            messagebox.showwarning("Validare", f"S-au găsit {issues} celule goale/invalide!")



    # ========================================================================

    # MINIMIZE TO TRAY

    # ========================================================================

    def minimize_to_tray(self):

        if not HAS_TRAY:

            messagebox.showerror("Eroare", "Biblioteca pystray nu este instalată. Instalați cu: pip install pystray pillow")

            return

        self.withdraw()

        image = Image.new('RGB', (64, 64), color='blue')

        draw = ImageDraw.Draw(image)

        draw.rectangle((16, 16, 48, 48), fill='white')

        icon = pystray.Icon("HRApp", image, "HR Generator", menu=pystray.Menu(

            pystray.MenuItem("Deschide", self.show_from_tray),

            pystray.MenuItem("Ieșire", self.quit_app)

        ))

        threading.Thread(target=icon.run, daemon=True).start()

        self.tray_icon = icon



    def show_from_tray(self):

        self.tray_icon.stop()

        self.deiconify()



    # ========================================================================

    # MACRO RECORDER

    # ========================================================================

    def start_recording(self):

        self.macro_recorder.start()

        self.log("Înregistrare macro pornită.")



    def stop_recording(self):

        code = self.macro_recorder.stop()

        self.log("Înregistrare macro oprită.")

        self.edit_macro(code)



    def edit_macro(self, code):

        win = tk.Toplevel(self)

        win.title("Editare macro")

        win.geometry("600x400")

        text = scrolledtext.ScrolledText(win, wrap='word')

        text.pack(fill='both', expand=True, padx=5, pady=5)

        text.insert('1.0', code)

        def save_macro():

            code = text.get('1.0', 'end-1c')

            with open('recorded_macro.py', 'w', encoding='utf-8') as f:

                f.write(code)

            messagebox.showinfo("Salvat", "Macro salvat în recorded_macro.py")

            win.destroy()

        ttk.Button(win, text="Salvează", command=save_macro).pack(pady=5)



    # ========================================================================

    # PLUGIN MARKETPLACE

    # ========================================================================

    def open_plugin_marketplace(self):

        webbrowser.open("https://github.com/topics/hr-plugin")

        url = tk.simpledialog.askstring("Plugin marketplace", "Introdu URL-ul pluginului (fișier .py):")

        if url:

            try:

                r = requests.get(url)

                if r.status_code == 200:

                    plugin_dir = Path('plugins')

                    plugin_dir.mkdir(exist_ok=True)

                    name = url.split('/')[-1]

                    with open(plugin_dir / name, 'wb') as f:

                        f.write(r.content)

                    self.log(f"Plugin descărcat: {name}")

                    plugin_manager.load_plugins()

            except Exception as e:

                messagebox.showerror("Eroare", f"Nu s-a putut descărca pluginul: {e}")



    # ========================================================================

    # SHAREPOINT

    # ========================================================================

    def connect_sharepoint(self):

        messagebox.showinfo("SharePoint", "Integrarea SharePoint necesită configurare manuală.\n"

                            "Instalați office365-rest-python-client și adăugați codul în plugin-uri.")



    # ========================================================================

    # ERROR HANDLING

    # ========================================================================

    def show_error(self, title, message, detail=None, offer_save=True):

        if offer_save:

            if messagebox.askyesno("Eroare", f"{message}\n\nDoriși să salvați un raport de eroare?"):

                self.save_bug_report(title, message, detail)

        messagebox.showerror(title, message)



    def save_bug_report(self, title, message, detail):

        try:

            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

            bug_dir = Path(f"bug_report_{timestamp}")

            bug_dir.mkdir()

            with open(bug_dir / "error.txt", 'w', encoding='utf-8') as f:

                f.write(f"Titlu: {title}\nMesaj: {message}\nDetalii: {detail or ''}\n")

                f.write(traceback.format_exc())

            audit_log = Path('hr_app_audit.log')

            if audit_log.exists():

                shutil.copy(audit_log, bug_dir / 'audit.log')

            config_file = Path('app_config.json')

            if config_file.exists():

                shutil.copy(config_file, bug_dir / 'config.json')

            with zipfile.ZipFile(f"bug_report_{timestamp}.zip", 'w') as zipf:

                for file in bug_dir.iterdir():

                    zipf.write(file, arcname=file.name)

            shutil.rmtree(bug_dir)

            self.log(f"Raport de eroare salvat: bug_report_{timestamp}.zip")

        except Exception as e:

            self.log(f"Eroare la salvarea raportului: {e}")



    # ========================================================================

    # UPDATE CHECK

    # ========================================================================

    def check_for_updates(self):

        self._check_updates(silent=True)



    def check_for_updates_now(self):

        self._check_updates(silent=False)



    def _check_updates(self, silent=False):

        version_url = "https://raw.githubusercontent.com/utilizator/repo/main/version.json"

        download_url_base = "https://github.com/utilizator/repo/releases/download/v{version}/app.zip"

        try:

            r = requests.get(version_url, timeout=5)

            if r.status_code == 200:

                data = r.json()

                remote_version = data.get('version', '0')

                download_url = data.get('download_url', download_url_base.format(version=remote_version))

                if remote_version > self.current_version:

                    if not silent:

                        msg = f"Există o versiune nouă: {remote_version}\nVersiunea curentă: {self.current_version}\n\nDoriși să descărcați și să instalați actualizarea?"

                        if messagebox.askyesno("Actualizare disponibilă", msg):

                            self._download_and_install(download_url, remote_version)

                    else:

                        self.log_queue.put(f"Versiune nouă disponibilă: {remote_version}. Accesați Ajutor Verifică actualizări pentru a instala.")

                else:

                    if not silent:

                        messagebox.showinfo("Actualizări", "Aveși cea mai recentă versiune.")

            else:

                if not silent:

                    messagebox.showerror("Eroare", "Nu s-a putut verifica actualizarea. Verificați conexiunea la internet.")

        except Exception as e:

            if not silent:

                messagebox.showerror("Eroare", f"Eroare la verificarea actualizării:\n{e}")



    def _download_and_install(self, url, new_version):

        try:

            self.log_queue.put("Încep descărcarea actualizării...")

            temp_dir = Path("update_temp")

            if temp_dir.exists():

                shutil.rmtree(temp_dir)

            temp_dir.mkdir()



            zip_path = temp_dir / "update.zip"

            r = requests.get(url, stream=True)

            with open(zip_path, 'wb') as f:

                for chunk in r.iter_content(chunk_size=8192):

                    f.write(chunk)



            with zipfile.ZipFile(zip_path, 'r') as zipf:

                zipf.extractall(temp_dir)



            new_files = list(temp_dir.glob("*.py"))

            if not new_files:

                raise Exception("Arhiva nu conține niciun fișier .py")

            new_main = new_files[0]



            current_file = Path(sys.argv[0]).resolve()

            backup_file = current_file.with_suffix('.bak')

            shutil.copy2(current_file, backup_file)



            shutil.copy2(new_main, current_file)



            shutil.rmtree(temp_dir)



            self.log_queue.put("Actualizare descărcată și instalată. Aplicașia se va reporni.")

            audit.log(action='update', details={'from': self.current_version, 'to': new_version})



            self.quit()

            subprocess.Popen([sys.executable, current_file])

            sys.exit()



        except Exception as e:

            messagebox.showerror("Eroare", f"Actualizarea a eșuat:\n{e}")



    # --------------------------------------------------------

    # Metode pentru drag-and-drop

    # --------------------------------------------------------

    def on_extract_drop(self, event):

        files = self.tk.splitlist(event.data)

        self.add_template_files(files)



    def on_render_drop(self, event):

        files = self.tk.splitlist(event.data)

        self.add_template_files(files)



    def on_scan_drop(self, event):

        files = self.tk.splitlist(event.data)

        self.add_template_files(files)



    # --------------------------------------------------------

    # Metode pentru tab-ul de extragere (existente)

    # --------------------------------------------------------

    def add_extract_files(self):

        files = filedialog.askopenfilenames(filetypes=[("Template files", "*.docx *.odt *.txt *.html")])

        if files:

            self.add_template_files(files)



    def remove_extract_files(self):

        self.remove_selected_files(self.extract_files_listbox)



    def on_extract_placeholder_double_click(self, event):

        selection = self.extract_placeholder_listbox.curselection()

        if not selection or not hasattr(self, 'placeholder_map'):

            return

        placeholder = self.extract_placeholder_listbox.get(selection[0])

        if placeholder in self.placeholder_map:

            file_path = self.placeholder_map[placeholder][0]

            self.open_in_word_viewer(file_path, placeholder)

        else:

            messagebox.showerror("Eroare", "Nu s-a găsit niciun fișier pentru acest placeholder.")



    def start_extract(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Adăugați cel pușin un fișier șablon.")

            return

        output_excel = self.excel_output_path.get().strip()

        if not output_excel:

            messagebox.showerror("Eroare", "Introduceși numele fișierului Excel de ieșire.")

            return



        self.extract_btn.configure(state='disabled')

        self.extract_only_btn.configure(state='disabled')

        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep extragerea placeholder-urilor și generarea Excel...")



        if HAS_ASYNC:

            thread = threading.Thread(target=self._extract_task_async, args=(self.template_files.copy(), output_excel, True))

        else:

            thread = threading.Thread(target=self._extract_task, args=(self.template_files.copy(), output_excel, True))

        thread.daemon = True

        thread.start()



    def start_extract_only(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Adăugați cel pușin un fișier șablon.")

            return



        self.extract_btn.configure(state='disabled')

        self.extract_only_btn.configure(state='disabled')

        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep extragerea placeholder-urilor (fără Excel)...")



        if HAS_ASYNC:

            thread = threading.Thread(target=self._extract_task_async, args=(self.template_files.copy(), None, False))

        else:

            thread = threading.Thread(target=self._extract_task, args=(self.template_files.copy(), None, False))

        thread.daemon = True

        thread.start()



    def _extract_task(self, files, output_excel, generate_excel):

        try:

            def progress_cb(current, total):

                self.after(0, lambda: self.set_progress(current, total))

            placeholders, placeholder_map = extract_all_placeholders_from_files(

                files, self.log_queue, progress_callback=progress_cb)

            self.placeholder_map = placeholder_map

            if not placeholders:

                self.log_queue.put("Nu s-au găsit placeholder-uri în șabloane.")

            else:

                self.log_queue.put(f"Placeholder-uri găsite: {placeholders}")

                if generate_excel and output_excel:

                    generate_excel_template(placeholders, output_excel, self.log_queue)

                self.after(0, lambda: self.display_extracted_placeholders(placeholders))

        except Exception as e:

            err_msg = str(e)

            err_tb = traceback.format_exc()

            self.after(0, lambda msg=err_msg, tb=err_tb: self.show_error("Eroare generare", msg, tb, offer_save=True))

        finally:

            self.log_queue.put("Procesul de extragere s-a încheiat.")

            self.after(0, lambda: self.extract_btn.configure(state='normal'))

            self.after(0, lambda: self.extract_only_btn.configure(state='normal'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: self.notify_done("Extragere finalizată!"))



    def _extract_task_async(self, files, output_excel, generate_excel):

        async def run():

            def progress_cb(current, total):

                self.after(0, lambda: self.set_progress(current, total))

            placeholders, placeholder_map = await extract_all_placeholders_async(files, progress_cb)

            self.placeholder_map = placeholder_map

            if not placeholders:

                self.log_queue.put("Nu s-au găsit placeholder-uri în șabloane.")

            else:

                self.log_queue.put(f"Placeholder-uri găsite: {placeholders}")

                if generate_excel and output_excel:

                    generate_excel_template(placeholders, output_excel, self.log_queue)

                self.after(0, lambda: self.display_extracted_placeholders(placeholders))

            self.log_queue.put("Procesul de extragere s-a încheiat.")

            self.after(0, lambda: self.extract_btn.configure(state='normal'))

            self.after(0, lambda: self.extract_only_btn.configure(state='normal'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: self.notify_done("Extragere finalizată!"))



        asyncio.run(run())



    def display_extracted_placeholders(self, placeholders):

        self.extract_placeholder_listbox.delete(0, tk.END)

        for ph in placeholders:

            self.extract_placeholder_listbox.insert(tk.END, ph)



    # --------------------------------------------------------

    # Metode pentru tab-ul de generare (extinse)

    # --------------------------------------------------------

    def add_render_files(self):

        files = filedialog.askopenfilenames(filetypes=[("Template files", "*.docx *.odt *.txt *.html")])

        if files:

            self.add_template_files(files)



    def remove_render_files(self):

        self.remove_selected_files(self.render_files_listbox)



    def browse_data_file(self):

        filename = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])

        if filename:

            self.data_file_path.set(filename)

            self.load_excel_columns()

            self.incarca_previzualizare_excel_async()

            self.load_global_excel_df(filename)



    def load_global_excel_df(self, filename):

        """Încărcăm asincron fișierul Excel pentru a fi disponibil global în alte tab-uri."""

        def loader():

            try:

                self.global_excel_df = pd.read_excel(filename, dtype=str).fillna('')

                self.global_excel_df.columns = self.global_excel_df.columns.str.strip()

                self.log_queue.put(('GLOBAL_EXCEL_LOADED', "Fișierul Excel a fost încărcat cu succes (Global)."))

            except Exception as e:

                self.log_queue.put(('EXCEL_ERROR', f"Nu s-a putut încărca Excel global: {e}"))

        threading.Thread(target=loader, daemon=True).start()



    def load_sheets(self):

        """Compatibilitate retroactivă – nu mai e nevoie de selecție manuală."""

        pass



    def load_excel_columns(self):

        """Încarcă coloanele din prima foaie a fișierului Excel."""

        file_path = self.data_file_path.get().strip()

        if not file_path or not os.path.exists(file_path):

            return

        try:

            df = pd.read_excel(file_path, sheet_name=0, nrows=0)

            columns = list(df.columns)



            # Actualizăm toate combo-urile de coloane disponibile

            combos_to_update = []

            if hasattr(self, 'folder_column_combo'): combos_to_update.append(self.folder_column_combo)

            if hasattr(self, 'email_column_combo'): combos_to_update.append(self.email_column_combo)

            if hasattr(self, 'subfolder_combo'): combos_to_update.append(self.subfolder_combo)



            for combo in combos_to_update:

                combo['values'] = columns



            if 'ID' in columns:

                self.folder_column.set('ID')

            elif columns:

                self.folder_column.set(columns[0])

        except Exception as e:

            self.log_queue.put(f"Eroare la încărcarea coloanelor: {e}")

            messagebox.showerror("Eroare", f"Nu s-au putut citi coloanele din Excel:\n{e}")





    def preview_document(self):

        """Previzualizare rând selectat în șablonul curent (Live Preview)"""

        try:

            # 1. Obșinem datele rândului selectat din Excel Tree

            if not hasattr(self, 'excel_tree'):

                show_toast(self, "⚠️ Deschideși tab-ul Vizualizare Excel!", type="info")

                return



            selection = self.excel_tree.selection()

            if not selection:

                show_toast(self, "⚠️ Selectați un rând din tabelul Excel!", type="info")

                return



            item = self.excel_tree.item(selection[0])

            row_values = item['values']

            cols = self.excel_tree['columns']

            row_data = dict(zip(cols, row_values))



            # 2. Șablonul

            if not self.template_files:

                show_toast(self, "⚠️ Nu există șabloane!", type="error")

                return

            template_path = getattr(self, 'current_preview_template', self.template_files[0])



            # 3. Randare

            temp_out = Path("tmp_preview.docx")

            context = {k: ('' if pd.isna(v) else v) for k, v in row_data.items()}



            from template_utils import (

                fuzzy_match_columns,

                extract_placeholders_from_file,

                render_document_from_template

            )



            all_ph = extract_placeholders_from_file(template_path)

            f_map = fuzzy_match_columns(all_ph, list(row_data.keys()))

            for ph, col in f_map.items():

                if ph not in context or str(context[ph]).strip() == '':

                    context[ph] = context.get(col, '')



            render_document_from_template(template_path, context, temp_out)



            # 4. Afisare

            self.load_word_preview_highlight(str(temp_out))

            show_toast(self, "✨ Previzualizare Live generată!")



        except Exception as e:

            self.show_error("Eroare Previzualizare", str(e))



    def preview_document_floating(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Nu există șabloane încărcate.")

            return

        data_file = self.data_file_path.get().strip()

        if not data_file or not os.path.exists(data_file):

            messagebox.showerror("Eroare", "Selectați mai întâi un fișier Excel.")

            return

        try:

            df = pd.read_excel(data_file, sheet_name=0)

        except Exception as e:

            messagebox.showerror("Eroare", f"Nu s-a putut citi fișierul Excel:\n{e}")

            return



        row_idx = self.preview_row_var.get() - 1

        if row_idx < 0 or row_idx >= len(df):

            messagebox.showerror("Eroare", "Numărul rândului este invalid.")

            return



        row_data = df.iloc[row_idx].to_dict()

        context = {k: ('' if pd.isna(v) else v) for k, v in row_data.items()}



        for script in self.scripts:

            try:

                context = script(context)

            except Exception as e:

                self.log_queue.put(f"Eroare în script la previzualizare: {e}")



        template_path = self.template_files[0]

        try:

            ext = Path(template_path).suffix

            preview_path = Path("temp_preview") / f"preview{ext}"

            preview_path.parent.mkdir(exist_ok=True)

            render_document_from_template(template_path, context, preview_path)



            win = tk.Toplevel(self)

            win.title(f"Previzualizare - Rând {self.preview_row_var.get()}")

            win.geometry("800x600")

            text_widget = scrolledtext.ScrolledText(win, wrap='word', font=('Courier', 10))

            text_widget.pack(fill='both', expand=True)



            if ext == '.docx':

                doc = Document(preview_path)

                content = '\n'.join([p.text for p in doc.paragraphs])

            elif ext == '.odt' and HAS_ODF:

                doc = load(preview_path)

                content = ''

                for elem in doc.getElementsByType(P):

                    if elem.firstChild and elem.firstChild.nodeType == elem.TEXT_NODE:

                        content += elem.firstChild.data + '\n'

            else:

                with open(preview_path, 'r', encoding='utf-8') as f:

                    content = f.read()



            text_widget.insert('1.0', content)

            for key, value in context.items():

                if value and str(value).strip():

                    start = 1.0

                    while True:

                        pos = text_widget.search(str(value), start, stopindex=tk.END, nocase=True)

                        if not pos:

                            break

                        end = f"{pos}+{len(str(value))}c"

                        text_widget.tag_add('highlight', pos, end)

                        start = end

            text_widget.tag_config('highlight', background='yellow')

            text_widget.configure(state='disabled')



        except Exception as e:

            messagebox.showerror("Eroare", f"Eroare la generarea previzualizării:\n{e}")



    def start_render(self):

        try:

            print(f"DEBUG: start_render entered | templates={len(self.template_files)}", flush=True)

            if getattr(self, 'is_generating', False):

                print("DEBUG: already generating", flush=True)

                return

            self.is_generating = True



            data_file = self.data_file_path.get().strip()

            output_dir = self.output_dir_path.get().strip()

            print(f"DEBUG: files ok | data={data_file} | out={output_dir}", flush=True)



            if not self.template_files:

                print("DEBUG: no template files, showing messagebox", flush=True)

                messagebox.showerror("Eroare", "Adăugați cel pușin un fișier șablon.")

                self.is_generating = False

                return

            if not data_file:

                print("DEBUG: no data file, showing messagebox", flush=True)

                messagebox.showerror("Eroare", "Selectați fișierul Excel completat.")

                self.is_generating = False

                return

            if not output_dir:

                print("DEBUG: no output dir, showing messagebox", flush=True)

                messagebox.showerror("Eroare", "Selectați dosarul de ieșire.")

                self.is_generating = False

                return



            # Prima foaie este selectată automat

            selected_sheets = [0]

            print("DEBUG: configuring buttons", flush=True)

            if hasattr(self, 'render_btn'): self.render_btn.configure(state='disabled')

            if hasattr(self, 'stop_button'): self.stop_button.configure(state='normal')



            self.stop_render_event.clear()

            print("DEBUG: calling log", flush=True)

            self.log("Încep generarea documentelor...")



            folder_column = self.folder_column.get().strip() if hasattr(self, 'folder_column') else ""

            resume = 0

            if hasattr(self, 'auto_recovery_var') and self.auto_recovery_var.get():

                resume = -1



            print(f"DEBUG: starting thread | col={folder_column}", flush=True)

            import threading

            thread = threading.Thread(target=self._render_task, args=(

                self.template_files.copy(), data_file, output_dir, folder_column, selected_sheets, resume))

            thread.daemon = True

            thread.start()

            print("DEBUG: thread started", flush=True)

        except Exception as e:

            print(f"DEBUG ERROR in start_render: {e}", flush=True)

            import traceback

            traceback.print_exc()

            self.is_generating = False



    def open_web_wizard(self):

        import webbrowser

        webbrowser.open("http://127.0.0.1:5000/")



    def _render_task(self, files, data_file, output_dir, folder_column, sheets, resume_flag):

        self.success_count = 0

        self.error_count = 0

        self.progress_val = 0

        try:

            resume_from = 0

            if resume_flag == -1:

                from pathlib import Path

                checkpoint_file = Path(output_dir) / "checkpoint.txt"

                if checkpoint_file.exists():

                    with open(checkpoint_file, 'r') as f:

                        try:

                            resume_from = int(f.read().strip())

                        except:

                            pass



            email_config_dict = None

            if hasattr(self, 'email_config') and self.email_config['enabled'].get():

                email_config_dict = {

                    'enabled': True,

                    'smtp_server': self.email_config['smtp_server'].get(),

                    'smtp_port': self.email_config['smtp_port'].get(),

                    'username': self.email_config['username'].get(),

                    'password': self.email_config['password'].get(),

                    'from': self.email_config['from'].get(),

                    'to': self.email_config['to'].get(),

                    'subject': self.email_config['subject'].get(),

                    'body': self.email_config['body'].get()

                }



            def progress_cb(current, total):

                self.after(0, lambda: self.set_progress(current, total))



            render_documents(

                files, data_file, output_dir, folder_column, sheets, self.log_queue, resume_from,

                progress_cb, self.stop_render_event, self.filename_pattern.get(),

                parallel=self.multiprocessing_var.get(),

                chunksize=self.chunksize.get(),

                email_config=email_config_dict,

                send_mode=self.email_send_mode.get(),

                email_column=self.email_column.get() if self.email_send_mode.get() != 'none' else None,

                email_subject_pattern=self.email_subject_pattern.get() if self.email_send_mode.get() != 'none' else None,

                pdf_gen=self.pdf_var.get(),

                zip_gen=self.zip_var.get(),

                merge_gen=self.merge_var.get(),

                audio_alert=self.audio_var.get(),

                subfolder_col=self.subfolder_var.get(),

                clean_data=self.clean_data_var.get(),

                zip_per_row=self.zip_per_row_var.get()

            )

        except Exception as e:

            self.log_queue.put(f"Eroare neșteptată: {e}")

            err_msg = str(e)

            err_tb = traceback.format_exc()

            self.after(0, lambda msg=err_msg, tb=err_tb: self.show_error("Eroare generare", msg, tb, offer_save=True))

        finally:

            self.is_generating = False

            self.after(0, lambda: self.render_btn.configure(state='normal'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: show_toast(self, "✅ Generare finalizată!"))



    def add_scan_files(self):

        files = filedialog.askopenfilenames(filetypes=[("Template files", "*.docx *.odt *.txt *.html")])

        if files:

            self.add_template_files(files)



    def remove_scan_files(self):

        self.remove_selected_files(self.scan_files_listbox)



    def start_scan_interactive(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Adăugați cel pușin un fișier pentru scanare.")

            return

        self.scan_interactive_btn.configure(state='disabled')

        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep scanarea interactivă...")

        thread = threading.Thread(target=self._scan_interactive_task, args=(self.template_files.copy(),))

        thread.daemon = True

        thread.start()



    def _scan_interactive_task(self, files):

        try:

            data = []

            for file_path in files:

                placeholders = extract_placeholders_from_file(file_path)

                for ph in placeholders:

                    issues = []

                    if ' ' in ph:

                        issues.append("spașiu")

                    if re.search(r'[ăâîșșĂÂÎȘț]', ph):

                        issues.append("diacritice")

                    if len(ph) > 40:

                        issues.append("lung >40")

                    issue_str = ", ".join(issues) if issues else "ok"

                    corrected = normalize_placeholder(ph)

                    data.append({

                        'placeholder': ph,

                        'fisier': Path(file_path).name,

                        'probleme': issue_str,

                        'corectat': corrected if issues else ph

                    })

            self.placeholder_data = data

            self.after(0, lambda: self.show_scan_results(data))

            self.log(f"Scanare interactivă finalizată. {len(data)} intrări.")

        except Exception as e:

            self.log_queue.put(f"Eroare la scanare: {e}")

        finally:

            self.after(0, lambda: self.scan_interactive_btn.configure(state='normal'))

            self.after(0, self.reset_progress)



    def show_scan_results(self, data):

        for widget in self.results_frame.winfo_children():

            widget.destroy()

        tree_frame = ttk.Frame(self.results_frame)

        tree_frame.grid(row=0, column=0, sticky='nsew')

        self.results_frame.grid_rowconfigure(0, weight=1)

        self.results_frame.grid_columnconfigure(0, weight=1)



        columns = ('select', 'placeholder', 'fisier', 'probleme', 'corectat')

        self.tree = ttk.Treeview(tree_frame, columns=columns, show='headings', height=15)

        self.tree.heading('select', text='Select')

        self.tree.heading('placeholder', text='Placeholder')

        self.tree.heading('fisier', text='Fișier')

        self.tree.heading('probleme', text='Probleme')

        self.tree.heading('corectat', text='Nume corectat')

        self.tree.column('select', width=50, anchor='center')

        self.tree.column('placeholder', width=200)

        self.tree.column('fisier', width=150)

        self.tree.column('corectat', width=200)



        vsb = ttk.Scrollbar(tree_frame, orient="vertical", command=self.tree.yview)

        hsb = ttk.Scrollbar(tree_frame, orient="horizontal", command=self.tree.xview)

        self.tree.configure(yscrollcommand=vsb.set, xscrollcommand=hsb.set)

        self.tree.grid(row=0, column=0, sticky='nsew')

        vsb.grid(row=0, column=1, sticky='ns')

        hsb.grid(row=1, column=0, sticky='ew')

        tree_frame.grid_rowconfigure(0, weight=1)

        tree_frame.grid_columnconfigure(0, weight=1)



        self.correction_vars = []

        self.placeholder_data = data

        for item in data:

            var = tk.BooleanVar(value=(item['probleme'] != "ok"))

            self.correction_vars.append(var)

            self.tree.insert('', 'end', values=(

                '☐',

                item['placeholder'],

                item['fisier'],

                item['probleme'],

                item['corectat']

            ), tags=('row',))



        def on_tree_click(event):

            item = self.tree.identify_row(event.y)

            if item:

                col = self.tree.identify_column(event.x)

                if col == '#1':

                    idx = self.tree.index(item)

                    self.correction_vars[idx].set(not self.correction_vars[idx].get())

                    self.tree.set(item, column='select', value='☑' if self.correction_vars[idx].get() else '☐')



        self.tree.bind('<Button-1>', on_tree_click)



        def on_double_click(event):

            item = self.tree.identify_row(event.y)

            if item:

                col = self.tree.identify_column(event.x)

                if col == '#5':

                    x, y, width, height = self.tree.bbox(item, column='corectat')

                    value = self.tree.item(item, 'values')[4]

                    entry = ttk.Entry(self.tree)

                    entry.place(x=x, y=y, width=width, height=height)

                    entry.insert(0, value)

                    entry.focus()

                    def save_edit(event=None):

                        new_value = entry.get()

                        entry.destroy()

                        self.tree.set(item, column='corectat', value=new_value)

                        idx = self.tree.index(item)

                        self.placeholder_data[idx]['corectat'] = new_value

                    entry.bind('<Return>', save_edit)

                    entry.bind('<FocusOut>', save_edit)

                else:

                    idx = self.tree.index(item)

                    file_name = self.placeholder_data[idx]['fisier']

                    placeholder = self.placeholder_data[idx]['placeholder']

                    file_path = None

                    for f in self.template_files:

                        if Path(f).name == file_name:

                            file_path = f

                            break

                    if file_path:

                        self.open_in_word_viewer(file_path, placeholder)

                    else:

                        messagebox.showerror("Eroare", "Fișierul nu a fost găsit.")



        self.tree.bind('<Double-1>', on_double_click)



        btn_frame = ttk.Frame(self.results_frame)

        btn_frame.grid(row=1, column=0, pady=5)

        ttk.Button(btn_frame, text="Rescanează", command=self.reset_scan_tab).pack(side='left', padx=5)

        ttk.Button(btn_frame, text="Aplică corecțiile selectate", command=self.apply_selected_corrections).pack(side='left', padx=5)

        ttk.Button(btn_frame, text="Generează raport text", command=self.start_scan_report).pack(side='left', padx=5)

        ttk.Button(btn_frame, text="Exportă raport HTML", command=self.export_scan_html).pack(side='left', padx=5)

        if HAS_REPORTLAB:

            ttk.Button(btn_frame, text="Exportă raport PDF", command=self.export_scan_pdf).pack(side='left', padx=5)



    def reset_scan_tab(self):

        for widget in self.frame_scan.winfo_children():

            widget.destroy()

        from tabs.scan_tab import setup_scan_tab

        setup_scan_tab(self, self.frame_scan)



    def open_in_word_viewer(self, file_path, placeholder):

        self.notebook.set("Vizualizare documente")

        self.word_viewer_load_file(file_path, placeholder)



    def apply_selected_corrections(self):

        if not hasattr(self, 'placeholder_data') or not self.placeholder_data:

            messagebox.showerror("Eroare", "Nu există date de corectat. Efectuați mai întâi o scanare.")

            return

        if not self.template_files:

            messagebox.showerror("Eroare", "Nu există fișiere șablon încărcate.")

            return



        if not messagebox.askyesno("Confirmare",

                                    "Această operașie va modifica fișierele originale.\n"

                                    "Se va crea un backup automat.\n"

                                    "Continuași?"):

            return



        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep aplicarea corecțiilor selectate...")



        thread = threading.Thread(target=self._apply_selected_task, args=(self.template_files.copy(),))

        thread.daemon = True

        thread.start()



    def _apply_selected_task(self, files):

        try:

            backups = []

            for file_path in files:

                file_path = Path(file_path)

                backup_folder = file_path.parent / "backup_interactiv"

                backup_folder.mkdir(exist_ok=True)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                backup_name = f"{file_path.stem}_{timestamp}{file_path.suffix}"

                backup_path = backup_folder / backup_name

                shutil.copy2(file_path, backup_path)

                backups.append((str(file_path), str(backup_path), timestamp))

            self.history.extend(backups)

            self.log_queue.put(f"Backup creat: {len(backups)} fișiere.")



            corrections = []

            for i, var in enumerate(self.correction_vars):

                if var.get():

                    old_ph = self.placeholder_data[i]['placeholder']

                    new_ph = self.placeholder_data[i]['corectat']

                    if old_ph != new_ph:

                        corrections.append((old_ph, new_ph))



            if not corrections:

                self.log_queue.put("Nicio corecție selectată.")

                return



            for file_path in files:

                file_path = Path(file_path)

                ext = file_path.suffix.lower()



                if ext == '.docx':

                    doc = Document(file_path)

                    modified = False



                    def replace_in_paragraph(paragraph):

                        nonlocal modified

                        text = paragraph.text

                        new_text = text

                        for old, new in corrections:

                            pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                            replacement = '{{' + new + '}}'

                            new_text, count = re.subn(pattern, replacement, new_text, flags=re.IGNORECASE)

                            if count > 0:

                                modified = True

                        if new_text != text:

                            paragraph.text = new_text



                    for paragraph in doc.paragraphs:

                        replace_in_paragraph(paragraph)

                    for table in doc.tables:

                        for row in table.rows:

                            for cell in row.cells:

                                for paragraph in cell.paragraphs:

                                    replace_in_paragraph(paragraph)

                    for section in doc.sections:

                        for header in [section.header, section.first_page_header, section.even_page_header]:

                            if header:

                                for paragraph in header.paragraphs:

                                    replace_in_paragraph(paragraph)

                        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:

                            if footer:

                                for paragraph in footer.paragraphs:

                                    replace_in_paragraph(paragraph)

                    if modified:

                        doc.save(file_path)

                        self.log_queue.put(f"Modificat: {file_path.name}")



                elif ext == '.odt' and HAS_ODF:

                    doc = load(file_path)

                    modified = False

                    for elem in doc.getElementsByType(P):

                        if elem.firstChild and elem.firstChild.nodeType == elem.TEXT_NODE:

                            text = elem.firstChild.data

                            new_text = text

                            for old, new in corrections:

                                pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                                replacement = '{{' + new + '}}'

                                new_text, count = re.subn(pattern, replacement, new_text, flags=re.IGNORECASE)

                                if count > 0:

                                    modified = True

                            if new_text != text:

                                elem.firstChild.data = new_text

                    if modified:

                        doc.save(file_path)

                        self.log_queue.put(f"Modificat: {file_path.name}")



                else:

                    with open(file_path, 'r', encoding='utf-8') as f:

                        content = f.read()

                    modified = False

                    for old, new in corrections:

                        old_tag = '{{' + old + '}}'

                        new_tag = '{{' + new + '}}'

                        pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                        replacement = '{{' + new + '}}'

                        new_content, count = re.subn(pattern, replacement, content, flags=re.IGNORECASE)

                        if count > 0:

                            modified = True

                            content = new_content

                    if modified:

                        with open(file_path, 'w', encoding='utf-8') as f:

                            f.write(content)

                        self.log_queue.put(f"Modificat: {file_path.name}")

                    else:

                        self.log_queue.put(f"  {file_path.name} nu a necesitat modificări.")



            self.log_queue.put("Corecțiile au fost aplicate.")

            audit.log(action='apply_corrections', details={'count': len(corrections)})



        except Exception as e:

            logger.error(f"Eroare la aplicarea corecțiilor: {e}\n{traceback.format_exc()}")

            self.log_queue.put(f"Eroare: {e}")

        finally:

            self.after(0, lambda: self.stop_button.configure(state='disabled'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: self.notify_done("Corecții aplicate!"))



    def start_scan_report(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Adăugați cel pușin un fișier pentru scanare.")

            return

        self.scan_report_btn.configure(state='disabled')

        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep generarea raportului text...")

        thread = threading.Thread(target=self._scan_report_task, args=(self.template_files.copy(),))

        thread.daemon = True

        thread.start()



    def _scan_report_task(self, files):

        try:

            def progress_cb(current, total):

                self.after(0, lambda: self.set_progress(current, total))

            report = scan_template_files(files, self.log_queue, progress_cb)

            self.last_scan_report = report

        except Exception as e:

            self.log_queue.put(f"Eroare: {e}")

        finally:

            self.after(0, lambda: self.scan_report_btn.configure(state='normal'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: self.notify_done("Raport generat!"))



    def export_scan_html(self):

        if not hasattr(self, 'last_scan_report') or not self.last_scan_report:

            messagebox.showerror("Eroare", "Nu există niciun raport generat. Rulați mai întâi o scanare.")

            return

        file_path = filedialog.asksaveasfilename(defaultextension=".html", filetypes=[("HTML files", "*.html")])

        if file_path:

            export_scan_report_html(self.last_scan_report, file_path)

            webbrowser.open(file_path)

            self.log_queue.put(f"Raport exportat: {file_path}")



    def export_scan_pdf(self):

        if not hasattr(self, 'last_scan_report') or not self.last_scan_report:

            messagebox.showerror("Eroare", "Nu există niciun raport generat. Rulați mai întâi o scanare.")

            return

        if not HAS_REPORTLAB:

            messagebox.showerror("Eroare", "ReportLab nu este instalat. Instalați cu: pip install reportlab")

            return

        file_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF files", "*.pdf")])

        if file_path:

            export_scan_report_pdf(self.last_scan_report, file_path)

            webbrowser.open(file_path)

            self.log_queue.put(f"Raport exportat: {file_path}")



    def start_correct(self):

        if not self.template_files:

            messagebox.showerror("Eroare", "Adăugați cel pușin un fișier pentru corectare.")

            return

        if not messagebox.askyesno("Confirmare",

                                    "Corecția automată va modifica fișierele originale.\n"

                                    "Se va crea un backup automat.\n"

                                    "Continuași?"):

            return

        self.correct_btn.configure(state='disabled')

        self.stop_button.configure(state='normal')

        self.stop_render_event.clear()

        self.log("Încep corecția automată...")

        thread = threading.Thread(target=self._correct_task, args=(self.template_files.copy(),))

        thread.daemon = True

        thread.start()



    def _correct_task(self, files):

        try:

            backups = []

            for file_path in files:

                file_path = Path(file_path)

                backup_folder = file_path.parent / "backup_interactiv"

                backup_folder.mkdir(exist_ok=True)

                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

                backup_name = f"{file_path.stem}_{timestamp}{file_path.suffix}"

                backup_path = backup_folder / backup_name

                shutil.copy2(file_path, backup_path)

                backups.append((str(file_path), str(backup_path), timestamp))

            self.history.extend(backups)

            self.log_queue.put(f"Backup creat: {len(backups)} fișiere.")

            self.after(0, lambda: self.history_listbox.insert(0, f"{len(backups)} backup-uri efectuate"))



            for file_path in files:

                file_path = Path(file_path)

                ext = file_path.suffix.lower()

                if ext == '.docx':

                    doc = Document(file_path)

                elif ext == '.odt' and HAS_ODF:

                    doc = load(file_path)

                else:

                    with open(file_path, 'r', encoding='utf-8') as f:

                        content = f.read()

                    modified = False

                    corrections = []

                    for i, var in enumerate(self.correction_vars):

                        if var.get():

                            old_ph = self.placeholder_data[i]['placeholder']

                            new_ph = self.placeholder_data[i]['corectat']

                            if old_ph != new_ph:

                                corrections.append((old_ph, new_ph))

                    for old, new in corrections:

                        pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                        replacement = '{{' + new + '}}'

                        content, count = re.subn(pattern, replacement, content, flags=re.IGNORECASE)

                        if count > 0:

                            modified = True

                    if modified:

                        with open(file_path, 'w', encoding='utf-8') as f:

                            f.write(content)

                        self.log_queue.put(f"Modificat: {file_path.name}")

                    else:

                        self.log_queue.put(f"  {file_path.name} nu a necesitat modificări.")

                    continue



                modified = False

                corrections = []

                for i, var in enumerate(self.correction_vars):

                    if var.get():

                        old_ph = self.placeholder_data[i]['placeholder']

                        new_ph = self.placeholder_data[i]['corectat']

                        if old_ph != new_ph:

                            corrections.append((old_ph, new_ph))



                def replace_in_paragraph(paragraph):

                    nonlocal modified

                    if not hasattr(paragraph, 'text'):

                        return

                    text = paragraph.text

                    new_text = text

                    for old, new in corrections:

                        pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                        replacement = '{{' + new + '}}'

                        new_text, count = re.subn(pattern, replacement, new_text, flags=re.IGNORECASE)

                        if count > 0:

                            modified = True

                    if new_text != text:

                        paragraph.text = new_text



                if ext == '.docx':

                    for paragraph in doc.paragraphs:

                        replace_in_paragraph(paragraph)

                    for table in doc.tables:

                        for row in table.rows:

                            for cell in row.cells:

                                for paragraph in cell.paragraphs:

                                    replace_in_paragraph(paragraph)

                    for section in doc.sections:

                        for header in [section.header, section.first_page_header, section.even_page_header]:

                            if header:

                                for paragraph in header.paragraphs:

                                    replace_in_paragraph(paragraph)

                        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:

                            if footer:

                                for paragraph in footer.paragraphs:

                                    replace_in_paragraph(paragraph)

                elif ext == '.odt' and HAS_ODF:

                    for elem in doc.getElementsByType(P):

                        if elem.firstChild and elem.firstChild.nodeType == elem.TEXT_NODE:

                            text = elem.firstChild.data

                            new_text = text

                            for old, new in corrections:

                                pattern = r'{{\s*' + re.escape(old) + r'\s*}}'

                                replacement = '{{' + new + '}}'

                                new_text, count = re.subn(pattern, replacement, new_text, flags=re.IGNORECASE)

                                if count > 0:

                                    modified = True

                            if new_text != text:

                                elem.firstChild.data = new_text



                if modified:

                    if ext == '.docx':

                        doc.save(file_path)

                    elif ext == '.odt' and HAS_ODF:

                        doc.save(file_path)

                    self.log_queue.put(f"Modificat: {file_path.name}")

                else:

                    self.log_queue.put(f"  {file_path.name} nu a necesitat modificări.")

            self.log_queue.put("Corecțiile au fost aplicate.")

            audit.log(action='apply_corrections', details={'count': len(corrections)})

        except Exception as e:

            self.log_queue.put(f"Eroare: {e}")

        finally:

            self.after(0, lambda: self.correct_btn.configure(state='normal'))

            self.after(0, self.reset_progress)

            self.after(0, lambda: self.notify_done("Corecții aplicate!"))



    def undo_last_correction(self):

        if not self.history:

            messagebox.showinfo("Info", "Nu există acțiuni de undo.")

            return

        self.show_history_dialog()



    def show_history_dialog(self):

        if not self.history:

            messagebox.showinfo("Istoric", "Nu există backup-uri.")

            return

        dialog = tk.Toplevel(self)

        dialog.title("Istoric corecții")

        dialog.geometry("700x400")

        dialog.transient(self)

        dialog.grab_set()



        tree = ttk.Treeview(dialog, columns=('file', 'backup', 'date'), show='headings')

        tree.heading('file', text='Fișier original')

        tree.heading('backup', text='Backup')

        tree.heading('date', text='Data')

        tree.column('file', width=250)

        tree.column('backup', width=250)

        tree.column('date', width=150)

        tree.pack(fill='both', expand=True, padx=10, pady=10)



        for file_path, backup_path, timestamp in self.history:

            tree.insert('', 'end', values=(file_path, backup_path, timestamp))



        def restore():

            selected = tree.selection()

            if not selected:

                return

            backup_path = tree.item(selected[0])['values'][1]

            file_path = tree.item(selected[0])['values'][0]

            if messagebox.askyesno("Confirmare", f"Restaurați {os.path.basename(file_path)} la această versiune?"):

                shutil.copy2(backup_path, file_path)

                self.log_queue.put(f"Restaurat {os.path.basename(file_path)} din backup.")

                audit.log(action='restore_backup', details={'file': file_path, 'backup': backup_path})

        ttk.Button(dialog, text="Restaurează", command=restore).pack(pady=5)



    # --------------------------------------------------------

    # Metode pentru tab-ul de vizualizare Word

    # --------------------------------------------------------

    def word_viewer_browse(self):

        file_path = filedialog.askopenfilename(filetypes=[("Word documents", "*.docx"), ("All files", "*.*")])

        if file_path:

            # Dacă fișierul este în output_dir, îl selectăm în combo

            out_dir = self.output_dir_path.get()

            if os.path.abspath(out_dir) in os.path.abspath(file_path):

                rel = os.path.relpath(file_path, out_dir)

                refresh_word_file_list(self)

                self.word_files_combo.set(rel)

            self.word_viewer_load_file(file_path)



    def word_viewer_load_file(self, file_path, highlight_placeholder=None, highlight_context=None):

        try:

            ext = Path(file_path).suffix.lower()

            if ext == '.docx':

                doc = Document(file_path)

                full_text = []

                for para in doc.paragraphs:

                    full_text.append(para.text)

                for table in doc.tables:

                    for row in table.rows:

                        for cell in row.cells:

                            for para in cell.paragraphs:

                                full_text.append(para.text)

                content = "\n".join(full_text)

            elif ext == '.odt' and HAS_ODF:

                doc = load(file_path)

                content = ''

                for elem in doc.getElementsByType(P):

                    if elem.firstChild and elem.firstChild.nodeType == elem.TEXT_NODE:

                        content += elem.firstChild.data + '\n'

            else:

                with open(file_path, 'r', encoding='utf-8') as f:

                    content = f.read()

            self.word_text.configure(state='normal')

            self.word_text.delete(1.0, tk.END)

            self.word_text.insert(1.0, content)



            # Highlight placeholders (new)

            highlight_placeholders_in_viewer(self)



            self.word_text.configure(state='disabled')

            self.word_current_file = file_path



            if highlight_placeholder:

                self.word_search_var.set(highlight_placeholder)

                self.word_viewer_search()

            elif highlight_context:

                self.word_text.configure(state='normal')

                self.word_text.tag_remove('highlight', 1.0, tk.END)

                for key, value in highlight_context.items():

                    if value and str(value).strip():

                        start = 1.0

                        while True:

                            pos = self.word_text.search(str(value), start, stopindex=tk.END, nocase=True)

                            if not pos:

                                break

                            end = f"{pos}+{len(str(value))}c"

                            self.word_text.tag_add('highlight', pos, end)

                            start = end

                self.word_text.tag_config('highlight', background='yellow')

                self.word_text.configure(state='disabled')

        except Exception as e:

            self.log_queue.put(f"Eroare la încărcarea fișierului: {e}")

            messagebox.showerror("Eroare", f"Nu s-a putut încărca fișierul:\n{e}")



    def word_viewer_search(self):

        search_term = self.word_search_var.get().strip()

        if not search_term:

            return

        self.word_text.configure(state='normal')

        self.word_text.tag_remove('search', 1.0, tk.END)

        start = 1.0

        self.word_search_results = []

        while True:

            pos = self.word_text.search(search_term, start, stopindex=tk.END, nocase=True)

            if not pos:

                break

            end = f"{pos}+{len(search_term)}c"

            self.word_text.tag_add('search', pos, end)

            self.word_search_results.append(pos)

            start = end

        self.word_text.tag_config('search', background='yellow')

        self.word_text.configure(state='disabled')



        if self.word_search_results:

            self.word_search_index = 0

            self.word_viewer_highlight_current()

            if hasattr(self, 'word_info_lbl'):

                self.word_info_lbl.configure(text=f"Găsite {len(self.word_search_results)} rezultate.")

        else:

            if hasattr(self, 'word_info_lbl'):

                self.word_info_lbl.configure(text="Nu s-au găsit rezultate.")

            else:

                messagebox.showinfo("Căutare", "Nu s-au găsit rezultate.")



    def word_viewer_search_next(self):

        if not self.word_search_results:

            return

        self.word_search_index = (self.word_search_index + 1) % len(self.word_search_results)

        self.word_viewer_highlight_current()



    def word_viewer_highlight_current(self):

        if not self.word_search_results:

            return

        pos = self.word_search_results[self.word_search_index]

        self.word_text.configure(state='normal')

        self.word_text.tag_remove('current', 1.0, tk.END)

        end = f"{pos}+{len(self.word_search_var.get())}c"

        self.word_text.tag_add('current', pos, end)

        self.word_text.tag_config('current', background='orange')

        self.word_text.see(pos)

        self.word_text.configure(state='disabled')



    def apply_preview(self):

        if not self.word_current_file or not hasattr(self, 'excel_data'):

            return

        idx = self.preview_row_combo.current()

        if idx < 0 or idx >= len(self.excel_data):

            return

        row_data = self.excel_data.iloc[idx].to_dict()

        row_data = {k: ('' if pd.isna(v) else v) for k, v in row_data.items()}

        ext = Path(self.word_current_file).suffix

        temp_file = f"temp_preview{ext}"

        render_document_from_template(self.word_current_file, row_data, temp_file)

        self.word_viewer_load_file(temp_file)



    # --------------------------------------------------------

    # Metode pentru tab-ul de vizualizare Excel (extinse)

    # --------------------------------------------------------

    def excel_viewer_browse(self):

        file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")])

        if file_path:

            self.excel_file_path.set(file_path)

            self.excel_viewer_load_sheet_list()



    def excel_viewer_load_from_render(self):

        file_path = self.data_file_path.get().strip()

        if file_path and os.path.exists(file_path):

            self.excel_file_path.set(file_path)

            self.excel_viewer_load_sheet_list()

        else:

            messagebox.showerror("Eroare", "Nu există un fișier Excel selectat în Pasul 2.")



    def excel_viewer_load_sheet_list(self):

        file_path = self.excel_file_path.get().strip()

        if not file_path or not os.path.exists(file_path):

            return

        try:

            xl = pd.ExcelFile(file_path)

            self.excel_sheet_combo['values'] = xl.sheet_names

            if xl.sheet_names:

                self.excel_sheet_var.set(xl.sheet_names[0])

                self.excel_viewer_load_sheet()

        except Exception as e:

            self.log_queue.put(f"Eroare la încărcarea foilor Excel: {e}")



    def excel_viewer_load_sheet(self):

        file_path = self.excel_file_path.get().strip()

        sheet = self.excel_sheet_var.get()

        if not file_path or not sheet:

            return

        try:

            self.excel_data = pd.read_excel(file_path, sheet_name=sheet)

            self.excel_filtered_data = self.excel_data.copy()

            self.excel_visible_columns = list(self.excel_data.columns)

            self.excel_viewer_refresh_columns_list()

            self.excel_stats_column['values'] = list(self.excel_data.columns)

            self.excel_viewer_display()

            self.preview_row_combo['values'] = list(range(1, len(self.excel_data)+1))

            self._populate_excel_tree(self.excel_data)

        except Exception as e:

            self.log_queue.put(f"Eroare la încărcarea foii: {e}")

            messagebox.showerror("Eroare", f"Nu s-au putut citi datele din Excel:\n{e}")



    def excel_viewer_refresh_columns_list(self):

        # Ne bazăm pe _populate_excel_tree care s-a apelat deja,

        # refacem checkbutton-urile folosind datele din excel_data

        if self.excel_data is not None:

            self._populate_excel_tree(self.excel_data)



    def excel_viewer_apply_columns(self):

        self.excel_visible_columns = [col for col, var in self.column_vars.items() if var.get()]

        self.excel_viewer_display()



    def excel_viewer_filter(self):

        if self.excel_data is None:

            return

        search_term = self.excel_search_var.get().strip()

        if not search_term:

            self.excel_filtered_data = self.excel_data.copy()

        else:

            mask = self.excel_data.apply(lambda row: row.astype(str).str.contains(search_term, case=False).any(), axis=1)

            self.excel_filtered_data = self.excel_data[mask].copy()

        self.excel_viewer_display()



    def excel_viewer_reset_filter(self):

        if self.excel_data is not None:

            self.excel_filtered_data = self.excel_data.copy()

            self.excel_search_var.set("")

            self.excel_viewer_display()



    def excel_viewer_display(self):

        for item in self.excel_tree.get_children():

            self.excel_tree.delete(item)

        if self.excel_filtered_data is None or len(self.excel_filtered_data) == 0:

            return

        cols = self.excel_visible_columns if self.excel_visible_columns else list(self.excel_filtered_data.columns)

        self.excel_tree['columns'] = cols

        for col in cols:

            self.excel_tree.heading(col, text=col, command=lambda c=col: self.excel_viewer_sort_by(c))

            self.excel_tree.column(col, width=120, minwidth=80)

        for idx, row in self.excel_filtered_data.iterrows():

            values = [row[col] if pd.notna(row[col]) else "" for col in cols]

            self.excel_tree.insert('', 'end', values=values, tags=(idx,))



    def excel_viewer_sort_by(self, col):

        if self.excel_filtered_data is None:

            return

        if hasattr(self, '_last_sort') and self._last_sort == col:

            ascending = not self._last_ascending

        else:

            ascending = True

        self._last_sort = col

        self._last_ascending = ascending

        self.excel_filtered_data = self.excel_filtered_data.sort_values(by=col, ascending=ascending).reset_index(drop=True)

        self.excel_viewer_display()



    def excel_viewer_generate_stats(self):

        col = self.excel_stats_column.get().strip()

        if not col or col not in self.excel_data.columns:

            messagebox.showerror("Eroare", "Selectați o coloană validă.")

            return

        stats = self.excel_data[col].value_counts(dropna=False)

        self.excel_stats_listbox.delete(0, tk.END)

        for val, cnt in stats.items():

            display_val = val if pd.notna(val) else "(gol)"

            self.excel_stats_listbox.insert(tk.END, f"{display_val}: {cnt}")



    def excel_viewer_stats_double_click(self, event):

        selection = self.excel_stats_listbox.curselection()

        if not selection:

            return

        line = self.excel_stats_listbox.get(selection[0])

        val = line.split(":")[0].strip()

        col = self.excel_stats_column.get().strip()

        if val == "(gol)":

            mask = self.excel_data[col].isna()

        else:

            mask = self.excel_data[col].astype(str) == val

        self.excel_filtered_data = self.excel_data[mask].copy()

        self.excel_viewer_display()



    def excel_viewer_export(self):

        if self.excel_filtered_data is None or len(self.excel_filtered_data) == 0:

            return

        cols = self.excel_visible_columns if self.excel_visible_columns else list(self.excel_filtered_data.columns)

        df_export = self.excel_filtered_data[cols].copy()

        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")])

        if file_path:

            if file_path.endswith('.csv'):

                df_export.to_csv(file_path, index=False)

            else:

                df_export.to_excel(file_path, index=False)

            messagebox.showinfo("Succes", f"Vizualizarea a fost exportată în {file_path}")



    # --------------------------------------------------------

    # Metode pentru tab-ul Bibliotecă (existente)

    # --------------------------------------------------------

    def load_library(self):

        conn = sqlite3.connect('template_library.db')

        c = conn.cursor()

        c.execute('''CREATE TABLE IF NOT EXISTS templates

                     (id INTEGER PRIMARY KEY AUTOINCREMENT,

                      name TEXT NOT NULL,

                      category TEXT,

                      file_path TEXT UNIQUE,

                      version TEXT,

                      author TEXT,

                      modified TIMESTAMP,

                      description TEXT)''')

        conn.commit()

        for row in self.library_tree.get_children():

            self.library_tree.delete(row)

        c.execute("SELECT id, name, category, version, author, modified FROM templates ORDER BY name")

        for row in c.fetchall():

            self.library_tree.insert('', 'end', values=row)

        conn.close()



    def import_to_library(self):

        files = filedialog.askopenfilenames(filetypes=[("Template files", "*.docx *.odt *.txt *.html")])

        if not files:

            return

        conn = sqlite3.connect('template_library.db')

        c = conn.cursor()

        for f in files:

            name = Path(f).name

            c.execute("SELECT id FROM templates WHERE file_path=?", (f,))

            if c.fetchone():

                continue

            modified = datetime.fromtimestamp(os.path.getmtime(f)).isoformat()

            c.execute("INSERT INTO templates (name, category, file_path, version, author, modified) VALUES (?, ?, ?, ?, ?, ?)",

                      (name, 'Necategorizat', f, '1.0', 'system', modified))

        conn.commit()

        conn.close()

        self.load_library()

        messagebox.showinfo("Succes", "Șabloane importate în bibliotecă.")



    def add_from_library(self):

        selected = self.library_tree.selection()

        if not selected:

            return

        item = self.library_tree.item(selected[0])

        template_id = item['values'][0]

        conn = sqlite3.connect('template_library.db')

        c = conn.cursor()

        c.execute("SELECT file_path FROM templates WHERE id=?", (template_id,))

        row = c.fetchone()

        conn.close()

        if row:

            self.add_template_files([row[0]])



    def delete_from_library(self):

        selected = self.library_tree.selection()

        if not selected:

            return

        if not messagebox.askyesno("Confirmare", "Ștergeși șablonul din bibliotecă?"):

            return

        item = self.library_tree.item(selected[0])

        template_id = item['values'][0]

        conn = sqlite3.connect('template_library.db')

        c = conn.cursor()

        c.execute("DELETE FROM templates WHERE id=?", (template_id,))

        conn.commit()

        conn.close()

        self.load_library()



    def backup_project(self):

        backup_name = f"backup_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"

        with zipfile.ZipFile(backup_name, 'w') as zipf:

            for f in self.template_files:

                zipf.write(f, arcname=Path(f).name)

            if self.data_file_path.get():

                zipf.write(self.data_file_path.get(), arcname="data.xlsx")

        self.log_queue.put(f"Backup creat: {backup_name}")

        audit.log(action='backup', details={'file': backup_name})



    def restore_backup(self):

        file_path = filedialog.askopenfilename(filetypes=[("ZIP files", "*.zip")])

        if not file_path:

            return

        with zipfile.ZipFile(file_path, 'r') as zipf:

            extract_dir = Path("restore_temp")

            extract_dir.mkdir(exist_ok=True)

            zipf.extractall(extract_dir)

            new_files = [str(extract_dir / f) for f in zipf.namelist() if f.endswith(('.docx', '.odt', '.txt', '.html'))]

            self.add_template_files(new_files)

            if 'data.xlsx' in zipf.namelist():

                self.data_file_path.set(str(extract_dir / 'data.xlsx'))

        self.log_queue.put(f"Proiect restaurat din {file_path}")

        audit.log(action='restore', details={'file': file_path})



    # --------------------------------------------------------

    # Metode pentru tab-ul Rapoarte (existente)

    # --------------------------------------------------------

    def generate_report(self):

        try:

            data = []

            if not os.path.exists('audit.jsonl'):

                self.log_queue.put("Fișierul audit.jsonl nu există. Generați mai întâi câteva documente.")

                return

            with open('audit.jsonl', 'r', encoding='utf-8') as f:

                for line in f:

                    record = json.loads(line)

                    if record['action'] == 'batch_render':

                        data.append(record)

            if not data:

                self.log_queue.put("Nu există date de raportare (nicio acțiune batch_render).")

                return

            dates = {}

            for rec in data:

                day = rec['timestamp'][:10]

                rows = rec['details'].get('rows', 0)

                dates[day] = dates.get(day, 0) + rows

            self.report_data = dates

            for widget in self.report_canvas_frame.winfo_children():

                widget.destroy()

            fig, ax = plt.subplots(figsize=(8,5))

            chart_type = self.chart_type.get()

            if chart_type == 'bar':

                ax.bar(dates.keys(), dates.values())

                ax.set_xlabel('Data')

                ax.set_ylabel('Număr documente generate')

                ax.set_title('Generări pe zi (bar)')

                plt.xticks(rotation=45)

            elif chart_type == 'pie':

                ax.pie(dates.values(), labels=dates.keys(), autopct='%1.1f%%')

                ax.set_title('Distribuție generări pe zi')

            elif chart_type == 'line':

                ax.plot(list(dates.keys()), list(dates.values()), marker='o')

                ax.set_xlabel('Data')

                ax.set_ylabel('Număr documente generate')

                ax.set_title('Generări pe zi (linie)')

                plt.xticks(rotation=45)

            elif chart_type == 'histogram':

                ax.hist(list(dates.values()), bins=5)

                ax.set_xlabel('Număr documente')

                ax.set_ylabel('Frecvenșă')

                ax.set_title('Histogramă generări')

            else:

                ax.bar(dates.keys(), dates.values())

                ax.set_title('Generări pe zi')

                plt.xticks(rotation=45)

            plt.tight_layout()

            canvas = FigureCanvasTkAgg(fig, master=self.report_canvas_frame)

            canvas.draw()

            canvas.get_tk_widget().pack(fill='both', expand=True)

            plt.close(fig)

        except Exception as e:

            self.log_queue.put(f"Eroare la generarea raportului: {e}")

            traceback.print_exc()

    def export_report_data(self):

        if not self.report_data:

            messagebox.showerror("Eroare", "Nu există date de raport. Generați mai întâi un raport.")

            return

        file_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV files", "*.csv")])

        if file_path:

            df = pd.DataFrame(list(self.report_data.items()), columns=['Data', 'Număr documente'])

            df.to_csv(file_path, index=False)

            self.log_queue.put(f"Date exportate în {file_path}")



    def export_report_excel(self):

        if not self.report_data:

            messagebox.showerror("Eroare", "Nu există date de raport. Generați mai întâi un raport.")

            return

        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])

        if file_path:

            df = pd.DataFrame(list(self.report_data.items()), columns=['Data', 'Număr documente'])

            df.to_excel(file_path, index=False)

            self.log_queue.put(f"Date exportate în {file_path}")

    # --------------------------------------------------------

    # Metode pentru tab-ul Rapoarte (noi)

    # --------------------------------------------------------

    def stats_post_categories(self):

        """Statistici pe categorii de post (grad profesional)."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        possible_cols = ['denumire_post', 'grad_profesional', 'Position', 'Categorie post', 'Grad profesional']

        col = None

        for c in possible_cols:

            if c in data.columns:

                col = c

                break

        if col is None:

            messagebox.showerror("Eroare", "Nu s-a găsit o coloană cu categoriile de post.")

            return

        counts = data[col].value_counts()

        self._show_bar_chart(counts, 'Categorii post', 'Număr persoane')



    def stats_departments(self):

        """Distribușia pe departamente/facultăși."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        possible_cols = ['Universitatea/departamentul', 'Department', 'Facultate', 'Departament']

        col = None

        for c in possible_cols:

            if c in data.columns:

                col = c

                break

        if col is None:

            messagebox.showerror("Eroare", "Nu s-a găsit o coloană cu departamente.")

            return

        counts = data[col].value_counts()

        self._show_bar_chart(counts, 'Departamente', 'Număr persoane')



    def stats_education_level(self):

        """Nivelul de educație (licenșă, master, doctorat)."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        possible_cols = ['Required_Education_Level', 'studii_de_specialitate', 'Education Level', 'Studii']

        col = None

        for c in possible_cols:

            if c in data.columns:

                col = c

                break

        if col is None:

            messagebox.showerror("Eroare", "Nu s-a găsit o coloană cu nivelul de educație.")

            return



        def classify(entry):

            if pd.isna(entry):

                return 'Necunoscut'

            text = str(entry).lower()

            if 'doctorat' in text or 'phd' in text:

                return 'Doctorat'

            elif 'master' in text:

                return 'Master'

            elif 'licenșă' in text or 'bachelor' in text or 'undergraduate' in text:

                return 'Licenșă'

            else:

                return 'Altele'



        levels = data[col].apply(classify)

        counts = levels.value_counts()

        self._show_pie_chart(counts, 'Nivel de educație')



    def stats_activities(self):

        """Repartizarea pe activităși (coloana cu numere separate)."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        possible_cols = ['Voi participa la activitatea/activitășile nr.:', 'Activităși', 'Activities']

        col = None

        for c in possible_cols:

            if c in data.columns:

                col = c

                break

        if col is None:

            messagebox.showerror("Eroare", "Nu s-a găsit o coloană cu activităși.")

            return

        from collections import Counter

        counter = Counter()

        for val in data[col].dropna():

            parts = str(val).split(';')

            for p in parts:

                p = p.strip()

                if p and p.isdigit():

                    counter[p] += 1

        if not counter:

            messagebox.showinfo("Info", "Nu există date despre activităși.")

            return

        series = pd.Series(counter).sort_index()

        self._show_bar_chart(series, 'Activitate', 'Număr persoane')



    def stats_id_types(self):

        """Statistici privind tipul actului de identitate."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        possible_cols = ['Tip act de identitate', 'Identity type', 'Document type', 'Tip BI/CI']

        col = None

        for c in possible_cols:

            if c in data.columns:

                col = c

                break

        if col is None:

            messagebox.showerror("Eroare", "Nu s-a găsit o coloană cu tipul actului de identitate.")

            return

        counts = data[col].value_counts()

        self._show_pie_chart(counts, 'Tip act de identitate')



    def skills_matrix(self):

        """Matricea de competenșe - top competenșe."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)



        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        skill_cols = ['Skills_Qualifications', 'Specific_Requirements', 'cunostinte_pc', 'cerinte_specifice', 'atributii_post', 'Competenșe']

        from collections import Counter

        counter = Counter()

        for col in skill_cols:

            if col in data.columns:

                for val in data[col].dropna():

                    import re

                    parts = re.split(r'[.,;/\n\r]+', str(val))

                    for p in parts:

                        p = p.strip()

                        if len(p) > 3:

                            counter[p] += 1

        if not counter:

            messagebox.showinfo("Info", "Nu există date despre competenșe.")

            return

        top = counter.most_common(20)

        labels = [item[0][:30] + ('...' if len(item[0])>30 else '') for item in top]

        values = [item[1] for item in top]

        self._show_bar_chart(pd.Series(values, index=labels), 'Competenșe', 'Frecvenșă')



    def generate_custom_chart(self):

        """Raport personalizat: utilizatorul alege coloanele și tipul de grafic."""

        data_source = getattr(self, 'global_excel_df', None)

        if data_source is None:

            data_source = getattr(self, 'excel_data', None)



        if data_source is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return

        x_col = self.custom_x_col.get()

        y_col = self.custom_y_col.get()

        chart_type = self.custom_chart_type.get()

        if not x_col or x_col == "(Niciunul)":

            messagebox.showerror("Eroare", "Selectați o coloană pentru axa X.")

            return

        # Filtrăm datele

        data = self.filtered_data if hasattr(self, 'filtered_data') and self.filtered_data is not None else data_source

        if chart_type == 'scatter':

            if not y_col:

                messagebox.showerror("Eroare", "Pentru scatter plot, selectați și o coloană Y.")

                return

            # Verificăm dacă ambele coloane sunt numerice

            x_num = pd.to_numeric(data[x_col], errors='coerce')

            y_num = pd.to_numeric(data[y_col], errors='coerce')

            if x_num.isna().all() or y_num.isna().all():

                messagebox.showerror("Eroare", "Coloanele selectate nu sunt numerice.")

                return

            self._show_scatter_plot(x_num, y_num, x_col, y_col)



        elif chart_type == 'bar':

            if y_col:

                # Bar chart cu două coloane: x ca etichete, y ca valori numerice

                y_num = pd.to_numeric(data[y_col], errors='coerce')

                if y_num.isna().all():

                    messagebox.showerror("Eroare", "Coloana Y nu este numerică.")

                    return

                # Grupăm după x și facem sumă folosind date numerice

                temp_df = data.copy()

                temp_df[y_col] = pd.to_numeric(temp_df[y_col], errors='coerce').fillna(0)

                grouped = temp_df.groupby(x_col)[y_col].sum()

                self._show_bar_chart(grouped, x_col, y_col)

            else:

                # Doar x: numărăm aparișiile

                counts = data[x_col].value_counts()

                self._show_bar_chart(counts, x_col, 'Număr aparișii')

        elif chart_type == 'pie':

            counts = data[x_col].value_counts()

            self._show_pie_chart(counts, x_col)

        elif chart_type == 'line':

            if not y_col:

                messagebox.showerror("Eroare", "Pentru line chart, selectați și o coloană Y.")

                return

            y_num = pd.to_numeric(data[y_col], errors='coerce')

            if y_num.isna().all():

                messagebox.showerror("Eroare", "Coloana Y nu este numerică.")

                return

            # Presupunem că x este ordonat (ex: date, index)

            fig, ax = plt.subplots(figsize=(10, 6))

            ax.plot(data[x_col], y_num, marker='o')

            ax.set_xlabel(x_col)

            ax.set_ylabel(y_col)

            ax.set_title(f'{y_col} în funcție de {x_col}')

            plt.xticks(rotation=45, ha='right')

            plt.tight_layout()

            self._display_figure(fig)

        else:

            messagebox.showerror("Eroare", "Tip de grafic necunoscut.")



    def _show_bar_chart(self, series, xlabel, ylabel):

        """Afișează un bar chart în report_canvas_frame."""

        if series.empty:

            self.log("Nu există date pentru acest grafic.")

            return

        fig, ax = plt.subplots(figsize=(12, 8), constrained_layout=True)

        series.plot(kind='bar', ax=ax, color='skyblue')

        ax.set_xlabel(xlabel)

        ax.set_ylabel(ylabel)

        ax.set_title(f'{ylabel} pe {xlabel}')

        plt.xticks(rotation=45, ha='right')

        self._display_figure(fig)



    def _show_pie_chart(self, series, title):

        """Afișează un pie chart."""

        if series.empty:

            self.log("Nu există date pentru acest grafic.")

            return

        fig, ax = plt.subplots(figsize=(10, 6))

        ax.pie(series, labels=series.index, autopct='%1.1f%%')

        ax.set_title(title)

        plt.tight_layout()

        self._display_figure(fig)



    def _show_scatter_plot(self, x_data, y_data, xlabel, ylabel):

        """Afișează un scatter plot."""

        if x_data.empty or y_data.empty:

            self.log("Nu există date pentru acest grafic.")

            return

        fig, ax = plt.subplots(figsize=(10, 6))

        ax.scatter(x_data, y_data, alpha=0.5)

        ax.set_xlabel(xlabel)

        ax.set_ylabel(ylabel)

        ax.set_title(f'Corelașie între {xlabel} și {ylabel}')

        plt.tight_layout()

        self._display_figure(fig)



    def _display_figure(self, fig):

        """Curășă frame-ul și afișează figura."""

        for widget in self.report_canvas_frame.winfo_children():

            widget.destroy()

        canvas = FigureCanvasTkAgg(fig, master=self.report_canvas_frame)

        canvas.draw()

        canvas.get_tk_widget().pack(fill='both', expand=True)

        plt.close(fig)



    def stats_gender(self):

        """Statistici pe gen (F/M) sau ghicit din Prenume."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)

        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return



        possible_cols = ['Sex', 'Gender', 'Gen']

        col = next((c for c in possible_cols if c in data.columns), None)



        if col:

            counts = data[col].value_counts()

        else:

            # Încercăm să ghicim din Prenume

            prenume_cols = ['Prenume', 'Nume mic', 'First Name', 'Prenume angajat']

            p_col = next((c for c in prenume_cols if c in data.columns), None)

            if p_col:

                def guess_gender(name):

                    if pd.isna(name): return "Necunoscut"

                    # Regula simplă: numele românești feminine se termină de obicei în 'a'

                    # Excepșii comune: Carmen, Beatrice, Iris etc.

                    name_clean = str(name).strip().split('-')[-1].split()[-1].lower()

                    if name_clean.endswith('a') or name_clean in ['carmen', 'iris', 'beatrice']:

                        return "Feminin"

                    return "Masculin"

                genders = data[p_col].apply(guess_gender)

                counts = genders.value_counts()

                self._show_pie_chart(counts, 'Distribuție pe gen (estimată din nume)')

                return

            else:

                messagebox.showerror("Eroare", "Nu s-a găsit coloană de Gen sau Prenume pentru estimare.")

                return



        self._show_pie_chart(counts, 'Distribuție pe gen')



    def stats_age_groups(self):

        """Statistici pe grupe de vârstă."""

        data = getattr(self, 'global_excel_df', None)

        if data is None:

            data = getattr(self, 'excel_data', None)

        if data is None:

            messagebox.showerror("Eroare", "Încărcați mai întâi un fișier Excel.")

            return



        # Încercăm să găsim coloana de vârstă

        age_col = None

        for c in ['Varsta', 'Age', 'Vârstă', 'vârstă']:

            if c in data.columns:

                age_col = c

                break



        if age_col:

            ages = pd.to_numeric(data[age_col], errors='coerce').dropna()

        else:

            # Încercăm din CNP dacă există

            cnp_col = None

            for c in ['CNP', 'cnp', 'Code']:

                if c in data.columns:

                    cnp_col = c

                    break

            if cnp_col:

                ages = []

                import datetime

                current_year = datetime.datetime.now().year

                for val in data[cnp_col].dropna():

                    s = str(val).strip()

                    if len(s) >= 7 and s.isdigit():

                        prefix = s[0]

                        year_part = int(s[1:3])

                        if prefix in '12': year = 1900 + year_part

                        elif prefix in '56': year = 2000 + year_part

                        else: year = 1900 + year_part

                        ages.append(current_year - year)

                ages = pd.Series(ages)

            else:

                messagebox.showerror("Eroare", "Nu s-a găsit nicio coloană pentru vârstă sau CNP.")

                return



        if ages.empty:

            messagebox.showinfo("Info", "Nu există date de vârstă valide.")

            return



        bins = [0, 18, 30, 45, 60, 100]

        labels = ['Sub 18', '18-30', '31-45', '46-60', 'Peste 60']

        groups = pd.cut(ages, bins=bins, labels=labels)

        counts = groups.value_counts().sort_index()

        self._show_bar_chart(counts, 'Grupă de vârstă', 'Număr persoane')



    def stats_template_usage(self):

        """Statistici privind utilizarea șabloanelor din audit.jsonl."""

        usage = {}

        try:

            import json

            with open('audit.jsonl', 'r', encoding='utf-8') as f:

                for line in f:

                    rec = json.loads(line)

                    if rec.get('action') == 'Render Document':

                        tmpl = rec.get('details', {}).get('file', 'Necunoscut')

                        usage[tmpl] = usage.get(tmpl, 0) + 1

        except Exception:

            pass



        if not usage:

            messagebox.showinfo("Info", "Nu există date de audit pentru șabloane.")

            return



        series = pd.Series(usage).sort_values(ascending=False)

        self._show_bar_chart(series, 'Șablon', 'Număr utilizări')



    def apply_filters(self):

        """Aplică filtrele definite de utilizator (placeholder)."""

        data_source = getattr(self, 'global_excel_df', None)

        if data_source is None:

            data_source = getattr(self, 'excel_data', None)

        if data_source is None:

            return

        self.filtered_data = data_source.copy()

        self.log("Filtre aplicate (implicit).")



    def export_filtered_data(self):

        """Exportă datele filtrate în CSV sau Excel."""

        data_to_export = getattr(self, 'filtered_data', None)

        if data_to_export is None:

            data_to_export = getattr(self, 'global_excel_df', None)

        if data_to_export is None:

            data_to_export = getattr(self, 'excel_data', None)



        if data_to_export is None:

            messagebox.showerror("Eroare", "Nu există date de exportat.")

            return



        file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",

                                                 filetypes=[("Excel files", "*.xlsx"), ("CSV files", "*.csv")])

        if file_path:

            try:

                if file_path.endswith('.csv'):

                    self.filtered_data.to_csv(file_path, index=False, encoding='utf-8-sig')

                else:

                    self.filtered_data.to_excel(file_path, index=False)

                self.log(f"Date exportate în {file_path}")

            except Exception as e:

                messagebox.showerror("Eroare", f"Nu s-a putut exporta: {e}")

    # --------------------------------------------------------

    # Metode pentru tab-ul Setări (existente)

    # --------------------------------------------------------

    def save_script(self):

        script_content = self.script_text.get(1.0, tk.END)

        with open('custom_script.py', 'w', encoding='utf-8') as f:

            f.write(script_content)

        try:

            spec = importlib.util.spec_from_file_location("custom_script", "custom_script.py")

            module = importlib.util.module_from_spec(spec)

            spec.loader.exec_module(module)

            if hasattr(module, 'process_row'):

                self.scripts = [module.process_row]

                self.log_queue.put("Script personalizat încărcat.")

                audit.log(action='script_loaded')

            else:

                messagebox.showerror("Eroare", "Scriptul trebuie să conțină funcția process_row(row).")

        except Exception as e:

            messagebox.showerror("Eroare", f"Nu s-a putut încărca scriptul:\n{e}")



    def load_script(self):

        if os.path.exists('custom_script.py'):

            with open('custom_script.py', 'r', encoding='utf-8') as f:

                self.script_text.delete(1.0, tk.END)

                self.script_text.insert(tk.END, f.read())

        else:

            messagebox.showinfo("Info", "Nu există un script salvat.")



    # --------------------------------------------------------

    # Metode pentru tab-ul Audit (existente)

    # --------------------------------------------------------

    def load_audit_log(self):

        for row in self.audit_tree.get_children():

            self.audit_tree.delete(row)

        try:

            with open('audit.jsonl', 'r', encoding='utf-8') as f:

                for line in f:

                    record = json.loads(line)

                    if self.audit_filter.get():

                        if self.audit_filter.get().lower() not in record['action'].lower():

                            continue



                    details = record.get('details', {})

                    if isinstance(details, dict):

                        detail_parts = []

                        if 'found_placeholders' in details or 'missing_placeholders' in details:

                            found = details.get('found_placeholders', [])

                            missing = details.get('missing_placeholders', [])

                            if found:

                                detail_parts.append(f"✅ Găsite: {', '.join(found)}")

                            if missing:

                                detail_parts.append(f"❌ Lipsă: {', '.join(missing)}")

                            for k, v in details.items():

                                if k not in ['found_placeholders', 'missing_placeholders']:

                                    detail_parts.append(f"{k}: {v}")

                            detail_str = " | ".join(detail_parts)

                        else:

                            detail_str = json.dumps(details, ensure_ascii=False)

                    else:

                        detail_str = str(details)



                    self.audit_tree.insert('', 'end', values=(

                        record.get('timestamp', ''),

                        record.get('user', ''),

                        record.get('action', ''),

                        detail_str

                    ))

        except FileNotFoundError:

            pass



    # --------------------------------------------------------

    # Metode pentru tab-ul Editor (existente)

    # --------------------------------------------------------

    def refresh_editor_list(self):

        self.editor_listbox.delete(0, tk.END)

        for f in self.template_files:

            self.editor_listbox.insert(tk.END, Path(f).name)



    def search_in_templates(self):

        term = self.search_var.get()

        if not term:

            return

        results = []

        for f in self.template_files:

            try:

                with open(f, 'r', encoding='utf-8', errors='ignore') as fp:

                    content = fp.read()

                    if term in content:

                        results.append(f)

            except:

                continue

        self.editor_text.delete(1.0, tk.END)

        self.editor_text.insert(1.0, "Rezultate căutare:\n" + "\n".join(results))



    def replace_in_templates(self):

        search = self.search_var.get()

        replace = self.replace_var.get()

        if not search:

            return

        count = 0

        for f in self.template_files:

            try:

                with open(f, 'r', encoding='utf-8', errors='ignore') as fp:

                    content = fp.read()

                if search in content:

                    new_content = content.replace(search, replace)

                    with open(f, 'w', encoding='utf-8') as fp:

                        fp.write(new_content)

                    count += 1

            except Exception as e:

                self.log(f"Eroare la procesarea {f}: {e}")

        self.log(f"Înlocuire finalizată în {count} fișiere.")



    # --------------------------------------------------------

    # Metode pentru tab-ul Statistici (existente)

    # --------------------------------------------------------

    def update_stats(self):

        """Actualizează tab-ul de statistici bazat pe datele din Excel-ul curent"""

        if not hasattr(self, 'stats_text'):

            return



        self.stats_text.configure(state='normal')

        self.stats_text.delete('1.0', tk.END)



        # Sursa de date: global_excel_df

        df = getattr(self, 'global_excel_df', None)

        if df is None:

            self.stats_text.insert(tk.END, "⚠️ Niciun fișier Excel global încărcat.\n")

            self.stats_text.insert(tk.END, "Incarcă un fișier Excel în tab-ul 'Randare' pentru a vedea statistici.\n")

            self.stats_text.configure(state='disabled')

            return



        total_rows = len(df)

        self.stats_text.insert(tk.END, f"SUMAR ANALITIC (Total: {total_rows} înregistrări)\n")

        self.stats_text.insert(tk.END, "="*60 + "\n\n")



        # 1. Distribuție pe Departamente

        dep_col = next((c for c in df.columns if 'departament' in str(c).lower() or 'universitate' in str(c).lower()), None)

        if dep_col:

            self.stats_text.insert(tk.END, f"DISTRIBUțIE PE DEPARTAMENTE:\n")

            counts = df[dep_col].value_counts().head(15)

            for name, count in counts.items():

                pct = (count / total_rows) * 100

                self.stats_text.insert(tk.END, f"  • {str(name):.<40} {count:>3} ({pct:>4.1f}%)\n")

            self.stats_text.insert(tk.END, "\n")



        # 2. Structură Posturi

        post_col = next((c for c in df.columns if 'post' in str(c).lower() or 'position' in str(c).lower()), None)

        if post_col:

            self.stats_text.insert(tk.END, f"TOP 10 POSTURI / POZIțII:\n")

            counts = df[post_col].value_counts().head(10)

            for name, count in counts.items():

                self.stats_text.insert(tk.END, f"  • {str(name):.<40} {count:>3}\n")

            self.stats_text.insert(tk.END, "\n")



        # 3. Nivel de Studii

        edu_col = next((c for c in df.columns if 'education' in str(c).lower() or 'studii' in str(c).lower()), None)

        if edu_col:

            self.stats_text.insert(tk.END, f"NIVEL EDUCAțIE:\n")

            counts = df[edu_col].value_counts()

            for name, count in counts.items():

                self.stats_text.insert(tk.END, f"  • {str(name):.<40} {count:>3}\n")

            self.stats_text.insert(tk.END, "\n")



        self.stats_text.insert(tk.END, f"\n✨ Analiza a fost actualizată la {datetime.now().strftime('%H:%M:%S')}\n")

        self.stats_text.configure(state='disabled')



    def export_audit_csv(self):

        try:

            data = []

            with open('audit.jsonl', 'r', encoding='utf-8') as f:

                for line in f:

                    data.append(json.loads(line))

            df = pd.DataFrame(data)

            file_path = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV", "*.csv")])

            if file_path:

                df.to_csv(file_path, index=False)

                self.log(f"Audit exportat în {file_path}")

        except Exception as e:

            self.show_error("Eroare", str(e))



    # --------------------------------------------------------

    # Asistent creare șabloane

    # --------------------------------------------------------

    def open_template_wizard(self):

        wizard = tk.Toplevel(self)

        wizard.title("Asistent creare șabloane")

        wizard.geometry("600x500")

        wizard.transient(self)

        wizard.grab_set()



        text = scrolledtext.ScrolledText(wizard, wrap='word', font=('Arial', 10))

        text.pack(fill='both', expand=True, padx=10, pady=10)

        text.insert('end',

            "Cum să creezi un șablon Word:\n"

            "1. Deschide un document Word existent sau creează unul nou.\n"

            "2. În locurile unde dorești să introduci date variabile, scrie numele placeholder-ului între acolade duble, de exemplu: {{nume}}.\n"

            "3. Poși folosi orice nume, dar evită spașiile și diacriticele (recomandat litere mici și underscore).\n"

            "4. Salvează documentul și încarcă-l în aplicație.\n\n"

            "Pentru a insera rapid un placeholder în documentul din Vizualizare Word, scrie numele mai jos și apasă 'Copiază'.\n"

        )

        text.configure(state='disabled')



        frame = ttk.Frame(wizard)

        frame.pack(fill='x', padx=10, pady=5)

        ttk.Label(frame, text="Nume placeholder:").pack(side='left')

        entry = ttk.Entry(frame, width=20)

        entry.pack(side='left', padx=5)

        def copy_placeholder():

            name = entry.get().strip()

            if name:

                self.clipboard_append('{{' + name + '}}')

                messagebox.showinfo("Copiat", "Textul a fost copiat în clipboard. Poși să-l lipești în document.")

        ttk.Button(frame, text="Copiază", command=copy_placeholder).pack(side='left')



    # --------------------------------------------------------

    # Funcșii comune

    # --------------------------------------------------------

    def browse_excel_output(self):

        filename = filedialog.asksaveasfilename(defaultextension=".xlsx", filetypes=[("Excel files", "*.xlsx")])

        if filename:

            self.excel_output_path.set(filename)



    def browse_output_dir(self):

        dirname = filedialog.askdirectory(title="Selectați dosarul de ieșire")

        if dirname:

            self.output_dir_path.set(dirname)



    def show_about(self):

        messagebox.showinfo("Despre", f"Generator documente din șabloane\nVersiune {self.current_version}\n\nRealizat de Sef Lucrari Moise Madalin Vasile\nProprietate personală\nUtilizați pe propria răspundere.")



    def stop_operation(self):

        self.stop_render_event.set()

        self.log("Comandă de oprire trimisă...")



    # --------------------------------------------------------

    # Inișializare tab-uri

    # --------------------------------------------------------

    def init_tabs(self):

        # Definim tab-urile în CTkTabview

        self.notebook.add("Dashboard")

        self.frame_dashboard = self.notebook.tab("Dashboard")

        setup_dashboard_tab(self, self.frame_dashboard)



        self.notebook.add("Pasul 1: Extrage placeholders")

        self.frame_extract = self.notebook.tab("Pasul 1: Extrage placeholders")

        setup_extract_tab(self, self.frame_extract)



        self.notebook.add("Pasul 2: Generează documente")

        self.frame_render = self.notebook.tab("Pasul 2: Generează documente")

        setup_render_tab(self, self.frame_render)



        self.notebook.add("Scanner șabloane")

        self.frame_scan = self.notebook.tab("Scanner șabloane")

        setup_scan_tab(self, self.frame_scan)



        self.notebook.add("Vizualizare documente")

        self.frame_word_viewer = self.notebook.tab("Vizualizare documente")

        setup_word_viewer_tab(self, self.frame_word_viewer)



        self.notebook.add("4. Vizualizare Excel")

        self.frame_excel_viewer = self.notebook.tab("4. Vizualizare Excel")

        setup_excel_viewer_tab(self, self.frame_excel_viewer)



        self.notebook.add("Bibliotecă șabloane")

        self.frame_library = self.notebook.tab("Bibliotecă șabloane")

        setup_library_tab(self, self.frame_library)



        self.notebook.add("Rapoarte")

        self.frame_reports = self.notebook.tab("Rapoarte")

        setup_reports_tab(self, self.frame_reports)



        self.notebook.add("Statistici")

        self.frame_stats = self.notebook.tab("Statistici")

        setup_stats_tab(self, self.frame_stats)



        self.notebook.add("Email")

        self.frame_email = self.notebook.tab("Email")

        setup_email_tab(self, self.frame_email)



        self.notebook.add("Setări")

        self.frame_settings = self.notebook.tab("Setări")

        setup_settings_tab(self, self.frame_settings)



        self.notebook.add("Audit Log")

        self.frame_audit = self.notebook.tab("Audit Log")

        setup_audit_tab(self, self.frame_audit)



        self.notebook.add("GHID INTELIGENT")

        self.frame_ai_guide = self.notebook.tab("GHID INTELIGENT")

        setup_ai_guide_tab(self, self.frame_ai_guide)



        self.notebook.add("MANUAL")

        self.frame_manual = self.notebook.tab("MANUAL")

        setup_manual_tab(self, self.frame_manual)



        # Setăm dashboard ca tab implicit

        self.notebook.set("Dashboard")



if __name__ == "__main__":

    # Multiprocessing init for Windows

    if sys.platform.startswith('win'):

        try:

            mp.set_start_method('spawn', force=True)

        except:

            pass

    # Scale text and UI elements up by 15% for better readability

    ctk.set_widget_scaling(1.15)

    ctk.set_window_scaling(1.15)



    app = PlaceholderApp()

    try:

        app.state("zoomed")  # Start maximized on Windows

    except:

        pass

    app.mainloop()


    def audit_empty_cells(self):
        """Identifica cine are celule goale."""
        data = self._get_stat_data()
        if data is None: return
        empty_rows = data[data.isna().any(axis=1) | (data == '').any(axis=1)]
        if empty_rows.empty:
            messagebox.showinfo("Audit", "Nu s-au gasit celule goale!")
            return
        name_col = self._find_col(data, ['Nume', 'nume_angajat', 'Nume de familie'])
        if name_col:
            names = empty_rows[name_col].head(20).tolist()
            msg = "Persoane cu date lipsa:\n" + "\n".join([f"- {n}" for n in names])
            if len(empty_rows) > 20:
                msg += f"\n... si inca {len(empty_rows)-20} rânduri."
            messagebox.showwarning("Audit Celule Goale", msg)
        else:
            messagebox.showwarning("Audit Celule Goale", f"Gasite {len(empty_rows)} rânduri cu date lipsa.")

    def audit_help_requests(self):
        """Identifica cine are nevoie de ajutor."""
        data = self._get_stat_data()
        if data is None: return
        possible_cols = ['ajutor', 'nevoie ajutor', 'asistenta']
        col = self._find_col(data, possible_cols)
        if not col:
            messagebox.showinfo("Audit", "Nu s-a gasit coloana de solicitare ajutor.")
            return
        help_requested = data[data[col].astype(str).lower().str.contains('da|yes|true', na=False)]
        if help_requested.empty:
            messagebox.showinfo("Audit", "Nimeni nu a solicitat ajutor.")
            return
        name_col = self._find_col(data, ['Nume', 'nume_angajat'])
        names = help_requested[name_col if name_col else data.columns[0]].head(20).tolist()
        msg = "Persoane care au solicitat ajutor:\n" + "\n".join([f"- {n}" for n in names])
        messagebox.showinfo("Solicitari Ajutor", msg)
